from __future__ import annotations

from pathlib import Path
from uuid import uuid4
from zipfile import ZipFile

import fitz
import pytest
from docx import Document
from markdown_it import MarkdownIt

from paperagent.agents.document_ir import (
    BlockKind,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    Provenance,
)
from paperagent.artifacts import ArtifactIntegrityError, ArtifactService
from paperagent.core.config import Settings
from paperagent.db import DatabaseManager
from paperagent.rendering.renderers import DocxRenderer, LatexRenderer, MarkdownRenderer
from paperagent.schemas.typography import TypographySpec

FIXTURE_ROOT = Path(__file__).resolve().parents[1] / "fixtures" / "document_layout"
PRIVATE_PLACEHOLDER = "Verified source content is supplied by the renderer."


def _semantic_document(tmp_path: Path) -> DocumentIR:
    image = tmp_path / "figure.png"
    canvas = fitz.open()
    page = canvas.new_page(width=480, height=180)
    page.draw_line((20, 90), (460, 90), color=(0, 0, 0))
    page.insert_text((24, 40), "synthetic figure")
    page.get_pixmap().save(image)
    canvas.close()
    return DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="Synthetic layout contract",
        language="mixed",
        sections=[
            DocumentSection(
                title="Experiment",
                goal="Lock semantic output contracts",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text="Structured body text.",
                        provenance=Provenance(agent="fixture"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.LIST,
                        text="first\nsecond",
                        provenance=Provenance(agent="fixture"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.TABLE,
                        caption="Table 1",
                        data={"rows": [["mode", "frequency"], ["1", "60"]]},
                        provenance=Provenance(agent="fixture"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.FIGURE,
                        caption="Figure 1",
                        data={"path": str(image)},
                        provenance=Provenance(agent="fixture"),
                    ),
                ],
            )
        ],
    )


def test_markdown_blob_is_parsed_and_private_front_matter_never_leaks(tmp_path: Path) -> None:
    source = (FIXTURE_ROOT / "semantic-report.md").read_text(encoding="utf-8")
    legacy = DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="Synthetic report",
        language="mixed",
        sections=[
            DocumentSection(
                title="正文",
                goal="Migration fixture",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text=source,
                        provenance=Provenance(agent="legacy-fixture"),
                    )
                ],
            )
        ],
    )
    output = MarkdownRenderer().render(legacy, tmp_path / "report.md")
    rendered = output.read_text(encoding="utf-8")
    tokens = MarkdownIt("commonmark").enable("table").parse(rendered)
    assert not rendered.startswith("---\n")
    assert "\\##" not in rendered and "\\**" not in rendered and "\\-" not in rendered
    assert sum(token.type == "heading_open" for token in tokens) >= 3
    assert sum(token.type == "table_open" for token in tokens) == 1


def test_docx_has_native_structure_media_headers_footers_and_page_number(
    tmp_path: Path,
) -> None:
    output = DocxRenderer().render(_semantic_document(tmp_path), tmp_path / "report.docx")
    word = Document(output)
    styles = [paragraph.style.name for paragraph in word.paragraphs]
    assert "Title" in styles and "Heading 1" in styles and "Caption" in styles
    assert len(word.tables) == 1
    assert len(word.inline_shapes) == 1
    captions = [
        paragraph.text
        for paragraph in word.paragraphs
        if paragraph.style.name == "Caption"
    ]
    assert captions == ["表 1", "图 1"]
    assert all(
        marker not in "\n".join(p.text for p in word.paragraphs)
        for marker in ("##", "**", "```", PRIVATE_PLACEHOLDER)
    )
    with ZipFile(output) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
        footer_xml = "".join(
            archive.read(name).decode("utf-8") for name in names if name.startswith("word/footer")
        )
        assert any(name.startswith("word/media/") for name in names)
        assert any(name.startswith("word/header") for name in names)
        assert any(name.startswith("word/footer") for name in names)
        assert "w:numPr" in document_xml
        assert "PAGE" in footer_xml


def test_xelatex_source_declares_complete_page_contract(tmp_path: Path) -> None:
    document = _semantic_document(tmp_path).restyle(
        TypographySpec(body_font="SimSun", body_size_pt=11)
    )
    source = LatexRenderer(executable="xelatex").source(document)
    assert PRIVATE_PLACEHOLDER not in source
    assert "a4paper" in source
    assert "\\usepackage{geometry}" in source
    assert "\\usepackage{fancyhdr}" in source
    assert "\\thepage" in source
    assert "\\includegraphics" in source
    assert "\\selectfont{}" in source
    assert "\\selectfont驻" not in source
    assert r"\renewcommand{\figurename}{图}" in source
    assert r"\renewcommand{\tablename}{表}" in source
    assert r"\caption{Table 1}" not in source
    assert r"\caption{Figure 1}" not in source


def test_reference_heading_is_localized_for_chinese_and_mixed_documents(
    tmp_path: Path,
) -> None:
    document = _semantic_document(tmp_path).model_copy(
        update={
            "metadata": {
                "evidence_manifest": [
                    {
                        "evidence_id": str(uuid4()),
                        "title": "来源",
                        "source_uri": "https://example.test/source",
                    }
                ]
            }
        }
    )
    source = LatexRenderer(executable="xelatex").source(document)
    assert r"\section*{参考文献}" in source
    assert r"\section*{References}" not in source
    output = DocxRenderer().render(document, tmp_path / "localized.docx")
    headings = [paragraph.text for paragraph in Document(output).paragraphs]
    assert "参考文献" in headings
    assert "References" not in headings


def _artifact_service(tmp_path: Path) -> tuple[ArtifactService, Path]:
    settings = Settings(
        project_root=tmp_path / "repo",
        data_dir=tmp_path / "data",
        environment="test",
    )
    databases = DatabaseManager(settings)
    databases.initialize_global()
    project_id = str(uuid4())
    root = databases.project_root(project_id)
    root.mkdir(parents=True)
    databases.project_engine(project_id).dispose()
    return ArtifactService(databases, project_id), root


def test_broken_or_zero_dimension_figure_cannot_become_valid(tmp_path: Path) -> None:
    service, root = _artifact_service(tmp_path)
    broken = root / "artifacts" / "broken.svg"
    broken.parent.mkdir(parents=True)
    broken.write_text("<svg><not-closed>", encoding="utf-8")
    with pytest.raises(ArtifactIntegrityError):
        service.register(
            broken,
            kind="figure",
            producer_tool="result.collect",
            run_id="run-figure",
        )


def test_style_only_revision_preserves_body_and_resource_semantics(tmp_path: Path) -> None:
    document = _semantic_document(tmp_path)
    before = document.model_dump(mode="json", exclude={"typography", "revision", "updated_at"})
    changed = document.restyle(TypographySpec(body_font="宋体", body_size_pt=11))
    after = changed.model_dump(mode="json", exclude={"typography", "revision", "updated_at"})
    assert before == after


def test_typography_workflow_contains_no_private_placeholder() -> None:
    source = (
        Path(__file__).resolve().parents[2]
        / "backend"
        / "paperagent"
        / "orchestration"
        / "interactive.py"
    ).read_text(encoding="utf-8")
    assert PRIVATE_PLACEHOLDER not in source
