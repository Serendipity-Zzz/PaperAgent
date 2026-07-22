import subprocess
from pathlib import Path
from uuid import uuid4

import fitz
from docx import Document
from markdown_it import MarkdownIt

from paperagent.agents.document_ir import (
    BlockKind,
    CitationRef,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    Provenance,
)
from paperagent.rendering.qa import PdfQualityAssurance
from paperagent.rendering.renderers import (
    DocxRenderer,
    LatexRenderer,
    MarkdownRenderer,
    TypstRenderer,
    WordPdfQueue,
)
from paperagent.rendering.template import DocxTemplateParser


def document(image: Path | None = None) -> DocumentIR:
    blocks = [
        DocumentBlock(
            kind=BlockKind.PARAGRAPH,
            text="中文 paragraph with [link](https://example.test).",
            provenance=Provenance(agent="test", author_viewpoint=True),
        ),
        DocumentBlock(
            kind=BlockKind.TABLE,
            data={"rows": [["列 A", "列 B"], ["1", "2"]]},
            caption="表 1",
            provenance=Provenance(agent="test"),
        ),
        DocumentBlock(
            kind=BlockKind.EQUATION,
            text="E = mc^2",
            provenance=Provenance(agent="test"),
        ),
        DocumentBlock(
            kind=BlockKind.CODE,
            text="print('safe')",
            data={"language": "python"},
            provenance=Provenance(agent="test"),
        ),
    ]
    if image:
        blocks.append(
            DocumentBlock(
                kind=BlockKind.FIGURE,
                caption="Figure 1",
                data={"path": str(image)},
                provenance=Provenance(agent="test"),
            )
        )
    return DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="PaperAgent 报告",
        language="mixed",
        sections=[DocumentSection(title="方法", goal="说明方法", blocks=blocks)],
    )


def test_docx_template_parser_never_uses_completed_body(tmp_path: Path) -> None:
    template = tmp_path / "学校模板.docx"
    word = Document()
    section = word.sections[0]
    section.header.paragraphs[0].text = "学校页眉"
    section.footer.paragraphs[0].text = "页脚"
    word.styles["Normal"].font.name = "宋体"
    word.add_heading("指定结构", level=1)
    word.add_paragraph("旧论文正文不能复用")
    word.save(template)
    profile = DocxTemplateParser().parse(template, completed_sample=True)
    assert profile.heading_structure == ["指定结构"]
    assert profile.completed_sample
    assert not profile.body_used_as_content
    assert any(region.kind == "header" for region in profile.regions)
    assert "旧论文正文不能复用" not in profile.model_dump_json()


def test_markdown_and_docx_renderers_roundtrip(tmp_path: Path) -> None:
    source = document()
    markdown = MarkdownRenderer().render(source, tmp_path / "paper.md")
    tokens = MarkdownIt().parse(markdown.read_text(encoding="utf-8"))
    assert any(token.type == "heading_open" for token in tokens)
    assert "E = mc^2" in markdown.read_text(encoding="utf-8")
    docx = DocxRenderer().render(source, tmp_path / "paper.docx")
    reopened = Document(docx)
    assert reopened.paragraphs[0].text == source.title
    assert len(reopened.tables) == 1


def test_typst_latex_word_detection_compile_and_no_overwrite(tmp_path: Path) -> None:
    source = document()

    def runner(command: list[str], cwd: Path, timeout: int) -> subprocess.CompletedProcess[str]:
        del timeout
        if "typst" in command[0]:
            Path(command[-1]).write_bytes(b"%PDF-mock")
        elif "xelatex" in command[0]:
            (cwd / "paper.pdf").write_bytes(b"%PDF-mock")
        elif "powershell" in command[0]:
            Path(command[-1]).write_bytes(b"%PDF-mock")
        return subprocess.CompletedProcess(command, 0, "ok", "")

    typst = TypstRenderer(executable="typst", runner=runner).render(source, tmp_path / "typst.pdf")
    assert typst.success and "Noto Serif CJK" in (tmp_path / "typst.typ").read_text("utf-8")
    latex = LatexRenderer(executable="xelatex", runner=runner).render(
        source, tmp_path / "paper.pdf"
    )
    assert latex.success and "xeCJK" in (tmp_path / "paper.tex").read_text("utf-8")
    missing = TypstRenderer(executable="", runner=runner)
    missing.executable = None
    assert missing.render(source, tmp_path / "missing.pdf").error_code == "TYPST_MISSING"
    docx = DocxRenderer().render(source, tmp_path / "word.docx")
    exported = WordPdfQueue(runner=runner).export(docx, tmp_path / "word.pdf")
    assert exported.success
    assert (
        WordPdfQueue(runner=runner).export(docx, tmp_path / "word.pdf").error_code
        == "OUTPUT_EXISTS"
    )


def test_pdf_page_render_and_visual_qa(tmp_path: Path) -> None:
    pdf = tmp_path / "qa.pdf"
    created = fitz.open()
    page = created.new_page()
    page.insert_text((72, 72), "Visible text")
    created.new_page()
    created.save(pdf)
    created.close()
    report = PdfQualityAssurance().inspect(pdf, tmp_path / "pages")
    assert report.page_count == 2
    assert len(report.rendered_pages) == 2
    assert all(Path(path).is_file() for path in report.rendered_pages)
    assert any(issue.code == "blank_or_scan" for issue in report.issues)


def test_figures_and_source_level_numeric_references_render_portably(tmp_path: Path) -> None:
    image = tmp_path / "chart.png"
    canvas = fitz.open()
    page = canvas.new_page(width=100, height=60)
    page.insert_text((10, 30), "chart")
    page.get_pixmap().save(image)
    canvas.close()
    source = document(image)
    first_id, second_id = uuid4(), uuid4()
    source.sections[0].blocks[0].citations = [
        CitationRef(evidence_id=first_id, verified=True),
        CitationRef(evidence_id=second_id, verified=True),
    ]
    source.metadata["evidence_manifest"] = [
        {
            "evidence_id": str(first_id),
            "title": "Shared paper metadata",
            "source_uri": "https://example.test/paper",
            "locator": {"page": 1},
        },
        {
            "evidence_id": str(second_id),
            "title": "Shared paper PDF",
            "source_uri": "https://example.test/paper",
            "locator": {"page": 2},
        },
    ]

    markdown = MarkdownRenderer().render(source, tmp_path / "portable.md").read_text("utf-8")
    typst = TypstRenderer().source(source)
    latex = LatexRenderer().source(source)
    docx = Document(DocxRenderer().render(source, tmp_path / "portable.docx"))

    assert markdown.count("- [1] Shared paper") == 1
    assert markdown.count(" [1]") >= 2
    assert str(first_id) not in markdown and str(second_id) not in markdown
    assert '#figure(image("' in typst and "caption: [Figure 1]" in typst
    assert "\\includegraphics[width=0.85\\linewidth," in latex
    assert "keepaspectratio" in latex
    assert len(docx.inline_shapes) == 1
    assert sum("[1] Shared paper" in paragraph.text for paragraph in docx.paragraphs) == 1
