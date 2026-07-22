from pathlib import Path
from uuid import uuid4

import fitz
import pytest
from docx import Document

from paperagent.agents.change_intent import ChangeIntent, ChangeIntentAgent, ChangeScope
from paperagent.agents.document_ir import (
    BlockKind,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    Provenance,
    diff_documents,
)
from paperagent.preview.schemas import Annotation, PreviewAnchor
from paperagent.rendering import (
    DocumentRevisionStore,
    DocxRenderer,
    FontResolver,
    LatexRenderer,
    PdfVisualDiff,
    RenderDependencyTracker,
    TargetedTypographyService,
    TypstRenderer,
)
from paperagent.rendering.artifacts import AnchorBinding


def local_document() -> DocumentIR:
    blocks = [
        DocumentBlock(
            kind=BlockKind.PARAGRAPH,
            text="正文不应被重写",
            provenance=Provenance(agent="test"),
        ),
        DocumentBlock(
            kind=BlockKind.TABLE,
            data={"rows": [["A", "B"], ["1", "2"]]},
            provenance=Provenance(agent="test"),
        ),
        DocumentBlock(
            kind=BlockKind.CODE,
            text="print('unchanged')",
            provenance=Provenance(agent="test"),
        ),
        DocumentBlock(
            kind=BlockKind.EQUATION,
            text="E=mc^2",
            provenance=Provenance(agent="test"),
        ),
    ]
    return DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="局部字体测试",
        language="mixed",
        sections=[DocumentSection(title="方法", goal="验证", blocks=blocks)],
    )


def test_block_patch_is_persistent_and_invalidates_only_one_logical_block(tmp_path: Path) -> None:
    original = local_document()
    target = original.sections[0].blocks[2]
    intent = ChangeIntent(
        scope=ChangeScope.BLOCK,
        block_ids=[target.block_id],
        typography_patch={"code_font": "Consolas", "code_size_pt": 9},
    )
    preview = ChangeIntentAgent.dry_run(original, intent)
    assert preview.content_preserved is True
    assert preview.before_hashes["content_hash"] == preview.after_hashes["content_hash"]
    assert preview.before_hashes["style_hash"] != preview.after_hashes["style_hash"]
    assert set(preview.affected_formats) == {"docx", "typst", "latex", "pdf"}
    updated, impact = ChangeIntentAgent.apply(original, intent)

    assert updated.sections == original.sections
    assert updated.resolve_typography(block_id=target.block_id).code_font == "Consolas"
    assert diff_documents(original, updated).changed_blocks == [target.block_id]
    assert impact.affected_blocks == [target.block_id]
    plan = RenderDependencyTracker().plan(original, updated, available_formats=["docx", "pdf"])
    assert plan.affected_block_ids == [target.block_id]
    assert not plan.content_changed
    assert not plan.regenerate_text and not plan.rerun_retrieval and not plan.rerun_experiments

    store = DocumentRevisionStore(tmp_path)
    store.save(original)
    store.save(updated)
    restored = store.load(updated.document_id)
    assert restored == updated


def test_section_patch_reaches_heading_table_code_formula_and_real_renderers(
    tmp_path: Path,
) -> None:
    original = local_document()
    section = original.sections[0]
    patch = {
        "heading_font": "Arial",
        "table_font": "Courier New",
        "code_font": "Consolas",
        "equation_font": "Cambria Math",
        "table_size_pt": 8,
        "code_size_pt": 9,
        "equation_size_pt": 10,
    }
    updated, impact = ChangeIntentAgent.apply(
        original,
        ChangeIntent(
            scope=ChangeScope.SECTION,
            section_ids=[section.section_id],
            typography_patch=patch,
        ),
    )
    assert set(impact.affected_blocks) == {item.block_id for item in section.blocks}

    path = DocxRenderer().render(updated, tmp_path / "local.docx")
    word = Document(path)
    assert word.paragraphs[1].runs[0].font.name == "Arial"
    assert word.tables[0].cell(0, 0).paragraphs[0].runs[0].font.name == "Courier New"
    code_paragraph = next(item for item in word.paragraphs if "print" in item.text)
    equation_paragraph = next(
        item for item in word.paragraphs if item.style.name == "PaperAgent Equation"
    )
    assert code_paragraph.runs[0].font.name == "Consolas"
    equation_xml = equation_paragraph._p.xml
    assert "Cambria Math" in equation_xml
    assert 'w:sz w:val="20"' in equation_xml
    assert "m:oMath" in equation_xml and "m:sSup" in equation_xml

    typst = TypstRenderer().source(updated)
    latex = LatexRenderer().source(updated)
    assert 'font: "Courier New"' in typst and 'font: "Consolas"' in typst
    assert "\\fontspec{Courier New}" in latex
    assert "\\setmathfont{Cambria Math}" in latex
    assert "\\mathversion" not in latex


@pytest.mark.anyio
async def test_preview_annotation_becomes_block_scoped_intent() -> None:
    document = local_document()
    block = document.sections[0].blocks[0]
    anchor = PreviewAnchor(
        source_file_id="file-1",
        source_hash="a" * 64,
        format="pdf",
        page=1,
        quote=block.text,
    )
    annotation = Annotation(
        project_id="project-1",
        artifact_id=uuid4(),
        anchor=anchor,
        body="正文字体设为宋体, 字号小四",
    )
    binding = AnchorBinding(
        artifact_id=annotation.artifact_id,
        anchor=anchor,
        section_id=document.sections[0].section_id,
        block_id=block.block_id,
    )
    intent = await ChangeIntentAgent().understand_annotation(annotation, binding)
    assert intent.scope is ChangeScope.BLOCK
    assert intent.block_ids == [block.block_id]
    assert intent.typography_patch["body_font"] == "宋体"


def test_pdf_visual_diff_persists_changed_page_evidence(tmp_path: Path) -> None:
    before = tmp_path / "before.pdf"
    after = tmp_path / "after.pdf"
    for path, text in ((before, "before"), (after, "after")):
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), text)
        document.save(path)
        document.close()

    report = PdfVisualDiff().compare(before, after, tmp_path / "diff")
    assert report.changed_pages == [1]
    assert report.pages[0].changed_pixel_ratio > 0
    assert Path(report.pages[0].diff_image or "").is_file()
    assert (tmp_path / "diff" / "visual-diff.json").is_file()


def test_targeted_service_renders_real_formats_and_records_dependency_scope(
    tmp_path: Path,
) -> None:
    font_dir = tmp_path / "fonts"
    font_dir.mkdir()
    (font_dir / "LocalCode.ttf").write_bytes(b"test font registry fixture")
    document = local_document()
    block = document.sections[0].blocks[2]
    result = TargetedTypographyService(
        tmp_path,
        fonts=FontResolver([font_dir]),
    ).apply(
        document,
        ChangeIntent(
            scope=ChangeScope.BLOCK,
            block_ids=[block.block_id],
            typography_patch={"code_font": "LocalCode", "code_size_pt": 10},
        ),
        formats=["md", "docx", "typst", "latex"],
    )

    assert result.invalidation.affected_block_ids == [block.block_id]
    assert not result.invalidation.regenerate_text
    assert {item.format for item in result.artifacts} == {"md", "docx", "typ", "tex"}
    assert all((tmp_path / item.path).is_file() for item in result.artifacts)
    assert DocumentRevisionStore(tmp_path).load(document.document_id) == result.document
