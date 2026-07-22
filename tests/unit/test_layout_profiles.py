from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from paperagent.rendering.fonts import FontResolver
from paperagent.rendering.layout import (
    ARCHETYPES,
    ArchetypeId,
    CascadeLayer,
    DocumentArchetypeClassifier,
    FontPlanService,
    FontRequest,
    FontRole,
    LengthUnit,
    LengthValue,
    NamedStyle,
    Orientation,
    PageCascade,
    PageOverride,
    PageSpec,
    StyleCascade,
    StyleOverride,
    StyleSheet,
    TocPolicy,
    archetype_layout_profile,
)
from paperagent.rendering.template import DocxTemplateParser, TemplateParseError


def test_all_ten_archetypes_have_complete_layout_and_qa_profiles() -> None:
    assert len(ARCHETYPES) == 10
    assert set(ARCHETYPES) == set(ArchetypeId)
    for archetype_id, archetype in ARCHETYPES.items():
        profile = archetype_layout_profile(archetype_id)
        assert archetype.required_sections
        assert archetype.default_components
        assert archetype.qa_rules
        assert profile.page.width.points() == pytest.approx(210 * 72 / 25.4)
        assert profile.page.height.points() == pytest.approx(297 * 72 / 25.4)
        assert set(profile.styles.styles) == set(NamedStyle)
    assert not ARCHETYPES[ArchetypeId.FORMAL_DOCUMENT].strict_official_standard


def test_archetype_classifier_prioritizes_user_then_structured_llm_then_fallback() -> None:
    classifier = DocumentArchetypeClassifier()
    explicit = classifier.classify(
        "ambiguous",
        explicit=ArchetypeId.MEETING_MINUTES,
    )
    assert explicit.archetype is ArchetypeId.MEETING_MINUTES
    assert explicit.confidence == 1 and not explicit.confirmation_required

    structured = classifier.classify(
        "ignored",
        structured={
            "archetype": "academic-paper",
            "confidence": 0.6,
            "evidence": ["abstract and references"],
            "alternatives": ["research-report"],
            "confirmation_required": False,
        },
    )
    assert structured.source == "llm" and structured.confirmation_required

    fallback = classifier.classify("请生成实验报告, 包含实验结果和实验图")
    assert fallback.archetype is ArchetypeId.EXPERIMENT_REPORT
    assert fallback.evidence


def test_page_spec_is_explicit_a4_and_units_are_serializable() -> None:
    page = PageSpec()
    assert page.orientation is Orientation.PORTRAIT
    assert page.content_width_pt > 0
    assert LengthValue(value=1, unit=LengthUnit.IN).points() == 72
    assert LengthValue(value=2, unit=LengthUnit.CM).points() == pytest.approx(2 * 72 / 2.54)
    assert LengthValue(value=2, unit=LengthUnit.CH).points(font_size_pt=11) == 22
    assert page.model_dump(mode="json")["width"] == {"value": 210.0, "unit": "mm"}

    resolved = PageCascade().resolve(
        [
            PageOverride(
                layer=CascadeLayer.TEMPLATE,
                source="uploaded-template",
                properties={"top_margin": {"value": 30, "unit": "mm"}},
            ),
            PageOverride(
                layer=CascadeLayer.TASK,
                source="user-current-task",
                properties={"top_margin": {"value": 20, "unit": "mm"}},
            ),
        ]
    )
    assert resolved.page.top_margin.points() == pytest.approx(20 * 72 / 25.4)
    assert resolved.sources["top_margin"] == "user-current-task"


def test_style_cascade_tracks_property_source_and_local_scope() -> None:
    overrides = [
        StyleOverride(
            layer=CascadeLayer.ARCHETYPE,
            source="experiment-report",
            style=NamedStyle.BODY_TEXT,
            properties={"font_size_pt": 10.5, "line_spacing": 1.25},
        ),
        StyleOverride(
            layer=CascadeLayer.PROJECT,
            source="project-preference",
            style=NamedStyle.BODY_TEXT,
            properties={"font_size_pt": 11},
        ),
        StyleOverride(
            layer=CascadeLayer.TASK,
            source="user-current-task",
            style=NamedStyle.BODY_TEXT,
            properties={"line_spacing": 1.5, "unknown_property": "diagnostic"},
        ),
        StyleOverride(
            layer=CascadeLayer.LOCAL,
            source="block-override",
            style=NamedStyle.BODY_TEXT,
            target_id="block-a",
            properties={"font_size_pt": 12},
        ),
    ]
    global_style = StyleCascade().resolve(NamedStyle.BODY_TEXT, overrides)
    local_style = StyleCascade().resolve(
        NamedStyle.BODY_TEXT,
        overrides,
        target_id="block-a",
    )
    assert global_style.properties.font_size_pt == 11
    assert global_style.properties.line_spacing == 1.5
    assert global_style.sources["line_spacing"] == "user-current-task"
    assert global_style.diagnostics[0].code == "UNSUPPORTED_STYLE_PROPERTY"
    assert local_style.properties.font_size_pt == 12
    assert local_style.sources["font_size_pt"] == "block-override"
    heading_one = StyleSheet().styles[NamedStyle.HEADING_1]
    assert heading_one.keep_with_next
    assert not heading_one.page_break_before


def test_toc_and_pagination_policies_are_deterministic() -> None:
    policy = TocPolicy()
    assert not policy.enabled(estimated_pages=5, section_count=3)
    assert policy.enabled(estimated_pages=6, section_count=3)
    assert policy.enabled(estimated_pages=3, section_count=4)
    assert not archetype_layout_profile(ArchetypeId.MEETING_MINUTES).toc.enabled(
        estimated_pages=20,
        section_count=20,
    )


def test_docx_template_profile_extracts_sections_styles_regions_and_capabilities(
    tmp_path: Path,
) -> None:
    path = tmp_path / "template.docx"
    word = Document()
    word.sections[0].header.paragraphs[0].text = "Template Header"
    word.sections[0].footer.paragraphs[0].text = "Template Footer"
    word.styles["Normal"].font.name = "Arial"
    word.styles["Normal"].font.size = Pt(11)
    word.add_heading("Expected structure", level=1)
    word.add_paragraph("Sample body must not become new content")
    word.add_table(rows=2, cols=2).style = "Table Grid"
    word.save(path)

    profile = DocxTemplateParser().parse(path, completed_sample=True)
    assert profile.section_count == 1
    assert profile.section_pages[0] == profile.page
    assert {item.kind for item in profile.regions} >= {"header", "footer"}
    assert profile.semantic_style_map["Normal"] == NamedStyle.BODY_TEXT.value
    assert "page" in profile.capabilities and "styles" in profile.capabilities
    assert profile.page_spec().width.points() == pytest.approx(profile.page.width_pt)
    assert any(item.style is NamedStyle.BODY_TEXT for item in profile.style_overrides())
    assert not profile.body_used_as_content
    assert "Sample body" not in profile.model_dump_json()

    with pytest.raises(TemplateParseError):
        DocxTemplateParser().parse(tmp_path / "missing.docx")


def test_missing_fonts_require_preview_and_installation_approval(tmp_path: Path) -> None:
    empty_fonts = tmp_path / "fonts"
    empty_fonts.mkdir()
    plan = FontPlanService(FontResolver([empty_fonts])).plan(
        [
            FontRequest(role=FontRole.EAST_ASIA, requested="Definitely Missing CJK"),
            FontRequest(role=FontRole.MATH, requested="Definitely Missing Math"),
        ]
    )
    assert plan.preview_required
    assert plan.installation_approval_required
    assert not plan.accepted
    assert all(item.requires_user_action for item in plan.resolutions)
