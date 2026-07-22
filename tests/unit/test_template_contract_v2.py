from __future__ import annotations

from pathlib import Path
from uuid import uuid4
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

from paperagent.agents.document_ir import (
    BlockKind,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    Provenance,
)
from paperagent.rendering import DocxRenderer
from paperagent.rendering.template import (
    DocxTemplateParser,
    TemplateApplicationMode,
    TemplateContractResolver,
    TemplateContractStore,
    TemplateNumberingContract,
    TemplateRepairService,
    migrate_template_contract,
)
from paperagent.schemas.numbering import NumberingOwner


def _template(path: Path, *, completed: bool = False) -> Path:
    word = Document()
    section = word.sections[0]
    section.header.paragraphs[0].text = "实验课程"
    section.footer.paragraphs[0].text = "第 PAGE 页"
    word.styles["Heading 1"].paragraph_format.keep_with_next = True
    word.add_heading("实验报告", level=1)
    word.add_paragraph("姓名: ________")
    word.add_paragraph("班级: ________")
    word.add_paragraph("指导教师: ________")
    word.add_paragraph("这是样例正文, 不得复制" if completed else "注意事项")
    word.add_paragraph("第一项", style="List Number")
    table = word.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    word.save(path)
    return path


def test_contract_v2_extracts_slots_styles_page_chrome_and_hashed_fixed_content(
    tmp_path: Path,
) -> None:
    path = _template(tmp_path / "template.docx", completed=True)
    parser = DocxTemplateParser()
    first = parser.parse_contract(path, completed_sample=True)
    second = parser.parse_contract(path, completed_sample=True)

    assert first == second
    assert first.schema_version == "2.0"
    assert first.source_hash == second.source_hash
    assert first.contract_hash == second.contract_hash
    assert {item.semantic_key for item in first.slots} >= {
        "author",
        "class_name",
        "advisor",
    }
    assert {item.kind for item in first.headers_footers} == {"header", "footer"}
    assert first.semantic_style_map["Heading 1"] == "Heading1"
    heading_style = next(item for item in first.styles if item.name == "Heading 1")
    assert heading_style.keep_with_next is True
    assert "Table Grid" in first.table_styles
    assert first.numbering.list_owned is True
    assert "这是样例正文, 不得复制" not in first.model_dump_json()
    assert all(item.category == "sample-content" for item in first.fixed_content)


def test_contract_store_is_content_addressed_and_roundtrips(tmp_path: Path) -> None:
    source = _template(tmp_path / "template.docx")
    parser = DocxTemplateParser()
    contract = parser.parse_contract(source)
    store = TemplateContractStore(tmp_path / "contracts")
    stored = store.save(contract)
    assert stored.name == f"{contract.source_hash}.json"
    assert store.load(contract.source_hash) == contract
    assert store.load("0" * 64) is None

    legacy = parser.parse(source).model_dump(mode="json")
    migrated = migrate_template_contract(legacy)
    assert migrated.schema_version == "2.0"
    assert migrated.fidelity_score <= 0.5
    assert migrated.diagnostics[0].code == "TEMPLATE_REEXTRACTION_RECOMMENDED"
    reextracted = migrate_template_contract(legacy, source_path=source)
    assert reextracted == contract


def test_active_and_external_package_parts_force_profile_only_mode(tmp_path: Path) -> None:
    path = _template(tmp_path / "unsafe.docx")
    with ZipFile(path, "a", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/vbaProject.bin", b"never execute")
        archive.writestr(
            "word/_rels/external-test.rels",
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="test" Target="https://example.test/x" '
            'TargetMode="External"/></Relationships>',
        )
        archive.writestr("word/embeddings/object1.bin", b"blocked")

    contract = DocxTemplateParser().parse_contract(path)
    codes = {item.code for item in contract.diagnostics}
    assert codes == {
        "TEMPLATE_MACRO_BLOCKED",
        "TEMPLATE_EXTERNAL_RELATIONSHIP_BLOCKED",
        "TEMPLATE_EMBEDDED_OBJECT_BLOCKED",
    }
    assert contract.safe_for_preserve is False
    decision = TemplateContractResolver().resolve(contract)
    assert decision.mode is TemplateApplicationMode.PROFILE_ONLY
    assert decision.numbering_owner is NumberingOwner.RENDERER

    document = DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="Safe output",
        language="en",
        sections=[
            DocumentSection(
                title="Result",
                goal="safe fallback",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text="Canonical content",
                        provenance=Provenance(agent="test"),
                    )
                ],
            )
        ],
    )
    output = DocxRenderer().render(document, tmp_path / "safe-output.docx", template=path)
    with ZipFile(output) as archive:
        names = set(archive.namelist())
    assert "word/vbaProject.bin" not in names
    assert not any(name.startswith("word/embeddings/") for name in names)


def test_safe_contract_resolver_explains_numbering_owner_and_mode(tmp_path: Path) -> None:
    contract = DocxTemplateParser().parse_contract(_template(tmp_path / "safe.docx"))
    decision = TemplateContractResolver().resolve(contract)
    assert decision.mode in {
        TemplateApplicationMode.PRESERVE,
        TemplateApplicationMode.REMAP,
    }
    expected_owner = (
        NumberingOwner.TEMPLATE
        if contract.numbering.heading_owned
        else NumberingOwner.RENDERER
    )
    assert decision.numbering_owner is expected_owner
    explicit = TemplateContractResolver().resolve(
        contract, explicit_mode=TemplateApplicationMode.PROFILE_ONLY
    )
    assert explicit.mode is TemplateApplicationMode.PROFILE_ONLY
    assert explicit.reason == "user-explicit"


def test_template_repair_changes_only_style_and_numbering_domains(tmp_path: Path) -> None:
    contract = DocxTemplateParser().parse_contract(_template(tmp_path / "repair.docx"))
    contract = contract.model_copy(
        update={
            "numbering": TemplateNumberingContract(
                owner=NumberingOwner.TEMPLATE,
                heading_owned=True,
                list_owned=True,
            )
        }
    )
    document = DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="Repair",
        language="en",
        sections=[
            DocumentSection(
                title="Body",
                goal="body",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text="Do not rewrite this content",
                        provenance=Provenance(agent="test"),
                    )
                ],
            )
        ],
    )
    updated, decision = TemplateRepairService().apply(document, contract)
    before = document.hashes()
    after = updated.hashes()
    assert decision.numbering_owner is NumberingOwner.TEMPLATE
    assert updated.revision == document.revision + 1
    assert after.style_hash != before.style_hash
    assert after.numbering_hash != before.numbering_hash
    assert after.content_hash == before.content_hash
    assert after.asset_set_hash == before.asset_set_hash
    assert after.citation_set_hash == before.citation_set_hash
