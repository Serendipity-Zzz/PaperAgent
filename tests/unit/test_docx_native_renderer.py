from __future__ import annotations

from pathlib import Path
from uuid import uuid4
from zipfile import ZipFile

import fitz
import pytest
from docx import Document

from paperagent.agents.document_ir import (
    BlockKind,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    EquationSpec,
    FigureSpec,
    InlineKind,
    InlineNode,
    ListItem,
    ListKind,
    ListSpec,
    Provenance,
    TableCell,
    TableRow,
    TableSpec,
)
from paperagent.rendering.renderers import DocxRenderer


def _document(tmp_path: Path) -> DocumentIR:
    image = tmp_path / "figure.png"
    drawing = fitz.open()
    page = drawing.new_page(width=480, height=180)
    page.insert_text((20, 40), "result figure")
    page.get_pixmap().save(image)
    drawing.close()
    return DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="Native Word 结构测试",
        language="mixed",
        metadata={"toc": True, "header_text": "实验报告", "footer_text": "内部文档"},
        sections=[
            DocumentSection(
                title="实验结果",
                goal="verify OOXML",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        inlines=[
                            InlineNode(kind=InlineKind.TEXT, text="访问 "),
                            InlineNode(
                                kind=InlineKind.LINK,
                                text="来源",
                                href="https://example.test/source",
                            ),
                            InlineNode(kind=InlineKind.STRONG, text=" 并复核结果"),
                            InlineNode(
                                kind=InlineKind.TEXT,
                                text=r", where $f_n=n\cdot f_1$.",
                            ),
                        ],
                        provenance=Provenance(agent="test"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.LIST,
                        list_spec=ListSpec(
                            kind=ListKind.ORDERED,
                            items=[
                                ListItem(text="prepare", children=[ListItem(text="nested")]),
                                ListItem(text="execute"),
                            ],
                        ),
                        provenance=Provenance(agent="test"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.TABLE,
                        caption="测量结果",
                        table=TableSpec(
                            rows=[
                                TableRow(
                                    cells=[
                                        TableCell(text="模态", header=True),
                                        TableCell(text=r"频率 $f_n$", header=True),
                                    ]
                                ),
                                TableRow(cells=[TableCell(text="1"), TableCell(text="60 Hz")]),
                            ]
                        ),
                        provenance=Provenance(agent="test"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.EQUATION,
                        text=r"y(x,t)=2A\sin(kx)\cos(\omega t)",
                        equation=EquationSpec(
                            latex=r"y(x,t)=2A\sin(kx)\cos(\omega t)",
                            number=True,
                        ),
                        provenance=Provenance(agent="test"),
                    ),
                    DocumentBlock(
                        kind=BlockKind.FIGURE,
                        caption="驻波实验结果",
                        figure=FigureSpec(path=str(image), alt_text="驻波振幅分布"),
                        provenance=Provenance(agent="test"),
                    ),
                ],
            )
        ],
    )


def test_docx_contains_a4_styles_fields_numbering_math_and_media(tmp_path: Path) -> None:
    output = DocxRenderer().render(_document(tmp_path), tmp_path / "report.docx")
    word = Document(output)
    assert word.sections[0].page_width.pt == pytest.approx(210 * 72 / 25.4, abs=0.1)
    assert word.sections[0].page_height.pt == pytest.approx(297 * 72 / 25.4, abs=0.1)
    assert {paragraph.style.style_id for paragraph in word.paragraphs} >= {
        "Title",
        "Heading1",
        "BodyText",
        "PaperAgentList",
        "Caption",
        "PaperAgentEquation",
    }
    assert len(word.tables) == 1
    assert len(word.inline_shapes) == 1
    figure_caption_index = next(
        index
        for index, paragraph in enumerate(word.paragraphs)
        if paragraph.style.style_id == "Caption" and "驻波实验结果" in paragraph.text
    )
    assert word.paragraphs[figure_caption_index - 1].paragraph_format.keep_with_next
    assert word.paragraphs[figure_caption_index].paragraph_format.keep_with_next is False

    with ZipFile(output) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        footer_xml = "".join(
            archive.read(name).decode("utf-8") for name in names if name.startswith("word/footer")
        )
        assert "w:pgSz" in document_xml and "w:pgMar" in document_xml
        assert "w:numPr" in document_xml
        assert "m:oMath" in document_xml
        assert "ω" in document_xml
        assert r"\omega" not in document_xml
        assert r"$f_n=n\cdot f_1$" not in document_xml
        assert r"$f_n$" not in document_xml
        assert "m:sSub" in document_xml
        assert "w:bookmarkStart" in document_xml
        assert "w:tblHeader" in document_xml and "w:cantSplit" in document_xml
        assert "w:tblGrid" in document_xml and 'w:type="dxa"' in document_xml
        assert "r:id" in document_xml and "hyperlink" in document_xml
        assert "BodyText" in styles_xml and "PaperAgentCode" in styles_xml
        assert "updateFields" in settings_xml
        assert "PAGE" in footer_xml and "NUMPAGES" in footer_xml
        assert any(name.startswith("word/media/") for name in names)
        assert any(name.startswith("word/header") for name in names)
        assert any(name.startswith("word/footer") for name in names)


def test_completed_template_preserves_brand_regions_but_not_sample_body(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    word = Document()
    word.sections[0].header.paragraphs[0].text = "School Header"
    word.sections[0].footer.paragraphs[0].text = "School Footer"
    word.add_heading("Old heading", level=1)
    word.add_paragraph("Old body must not leak")
    word.save(template)

    output = DocxRenderer().render(
        _document(tmp_path),
        tmp_path / "from-template.docx",
        template=template,
    )
    rendered = Document(output)
    body = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Old body must not leak" not in body
    assert "Native Word 结构测试" in body
    assert rendered.sections[0].header.paragraphs[0].text == "School Header"
    assert rendered.sections[0].footer.paragraphs[0].text == "School Footer"
