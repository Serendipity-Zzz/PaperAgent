from __future__ import annotations

import hashlib
import json
import re
from enum import StrEnum
from pathlib import Path
from typing import TYPE_CHECKING
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

from docx import Document
from pydantic import BaseModel, ConfigDict, Field

from paperagent.rendering.layout import (
    CascadeLayer,
    LengthUnit,
    LengthValue,
    NamedStyle,
    PageSpec,
    StyleOverride,
    semantic_style_id,
)
from paperagent.schemas.numbering import NumberingOwner

if TYPE_CHECKING:
    from paperagent.agents.document_ir import DocumentIR


class PageProfile(BaseModel):
    width_pt: float | None
    height_pt: float | None
    top_margin_pt: float | None
    bottom_margin_pt: float | None
    left_margin_pt: float | None
    right_margin_pt: float | None


class StyleProfile(BaseModel):
    name: str
    style_type: str
    font_name: str | None = None
    east_asia_font: str | None = None
    ascii_font: str | None = None
    hansi_font: str | None = None
    font_size_pt: float | None = None
    bold: bool | None = None
    italic: bool | None = None
    alignment: str | None = None
    line_spacing: float | None = None
    space_before_pt: float | None = None
    space_after_pt: float | None = None
    keep_with_next: bool | None = None
    keep_together: bool | None = None
    page_break_before: bool | None = None
    based_on: str | None = None


class TemplateRegion(BaseModel):
    kind: str
    text: str
    section_index: int | None = None


class TemplateProfile(BaseModel):
    source_name: str
    page: PageProfile
    section_pages: list[PageProfile] = Field(default_factory=list)
    section_count: int
    styles: list[StyleProfile]
    regions: list[TemplateRegion]
    heading_structure: list[str]
    has_numbering: bool
    numbering_abstract_count: int = 0
    has_toc: bool = False
    table_styles: list[str] = Field(default_factory=list)
    theme_fonts: list[str] = Field(default_factory=list)
    semantic_style_map: dict[str, str] = Field(default_factory=dict)
    capabilities: list[str] = Field(default_factory=list)
    missing_capabilities: list[str] = Field(default_factory=list)
    completed_sample: bool
    body_used_as_content: bool = False
    warnings: list[str] = Field(default_factory=list)

    def page_spec(self) -> PageSpec:
        default = PageSpec()

        def length(value: float | None, fallback: LengthValue) -> LengthValue:
            return LengthValue(value=value, unit=LengthUnit.PT) if value is not None else fallback

        return PageSpec(
            width=length(self.page.width_pt, default.width),
            height=length(self.page.height_pt, default.height),
            top_margin=length(self.page.top_margin_pt, default.top_margin),
            bottom_margin=length(self.page.bottom_margin_pt, default.bottom_margin),
            left_margin=length(self.page.left_margin_pt, default.left_margin),
            right_margin=length(self.page.right_margin_pt, default.right_margin),
        )

    def style_overrides(self) -> list[StyleOverride]:
        overrides: list[StyleOverride] = []
        for item in self.styles:
            semantic_value = self.semantic_style_map.get(item.name)
            if semantic_value is None:
                continue
            properties: dict[str, object] = {}
            if item.font_name:
                properties["font_family"] = item.font_name
            if item.east_asia_font:
                properties["east_asia_font"] = item.east_asia_font
            if item.font_size_pt:
                properties["font_size_pt"] = item.font_size_pt
            for key in (
                "bold",
                "italic",
                "line_spacing",
                "space_before_pt",
                "space_after_pt",
                "keep_with_next",
                "keep_together",
                "page_break_before",
            ):
                value = getattr(item, key)
                if value is not None:
                    properties[key] = value
            if properties:
                overrides.append(
                    StyleOverride(
                        layer=CascadeLayer.TEMPLATE,
                        source=f"template:{self.source_name}:{item.name}",
                        style=NamedStyle(semantic_value),
                        properties=properties,
                    )
                )
        return overrides


class TemplateDiagnosticSeverity(StrEnum):
    INFO = "info"
    WARNING = "warning"
    ERROR = "error"


class TemplateDiagnostic(BaseModel):
    model_config = ConfigDict(extra="forbid")

    code: str
    message: str
    severity: TemplateDiagnosticSeverity
    repair_node: str = "document.template.inspect"


class TemplateSlot(BaseModel):
    model_config = ConfigDict(extra="forbid")

    slot_id: str
    semantic_key: str
    label: str
    location: str
    required: bool = False
    source_style: str | None = None


class TemplateFixedContent(BaseModel):
    model_config = ConfigDict(extra="forbid")

    node_id: str
    text_hash: str
    category: str
    location: str
    source_style: str | None = None
    reusable: bool = False


class TemplateNumberingLevel(BaseModel):
    model_config = ConfigDict(extra="forbid")

    abstract_id: str
    level: int = Field(ge=0, le=8)
    format: str
    level_text: str
    start: int = Field(default=1, ge=0)
    restart: int | None = None
    paragraph_style: str | None = None


class TemplateNumberingContract(BaseModel):
    model_config = ConfigDict(extra="forbid")

    owner: NumberingOwner = NumberingOwner.NONE
    heading_owned: bool = False
    list_owned: bool = False
    levels: list[TemplateNumberingLevel] = Field(default_factory=list)
    style_links: dict[str, str] = Field(default_factory=dict)
    diagnostics: list[TemplateDiagnostic] = Field(default_factory=list)


class TemplateContractV2(BaseModel):
    model_config = ConfigDict(extra="forbid")

    source_hash: str = Field(pattern=r"^[0-9a-f]{64}$")
    package_type: str = "docx"
    page_sections: list[PageProfile]
    styles: list[StyleProfile]
    semantic_style_map: dict[str, str]
    numbering: TemplateNumberingContract
    theme_fonts: list[str]
    headers_footers: list[TemplateRegion]
    table_styles: list[str]
    slots: list[TemplateSlot]
    fixed_content: list[TemplateFixedContent]
    capabilities: list[str]
    fidelity_score: float = Field(ge=0, le=1)
    diagnostics: list[TemplateDiagnostic] = Field(default_factory=list)
    schema_version: str = "2.0"

    @property
    def contract_hash(self) -> str:
        payload = json.dumps(
            self.model_dump(mode="json"),
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        )
        return hashlib.sha256(payload.encode("utf-8")).hexdigest()

    @property
    def safe_for_preserve(self) -> bool:
        return not any(
            item.severity is TemplateDiagnosticSeverity.ERROR for item in self.diagnostics
        )


class TemplateApplicationMode(StrEnum):
    PRESERVE = "preserve"
    REMAP = "remap"
    PROFILE_ONLY = "profile-only"
    CLARIFY = "clarify"


class TemplateApplicationDecision(BaseModel):
    mode: TemplateApplicationMode
    reason: str
    confirmation_required: bool = False
    numbering_owner: NumberingOwner
    fidelity_score: float = Field(ge=0, le=1)
    diagnostics: list[TemplateDiagnostic] = Field(default_factory=list)


class TemplateContractResolver:
    def resolve(
        self,
        contract: TemplateContractV2,
        *,
        explicit_mode: TemplateApplicationMode | None = None,
    ) -> TemplateApplicationDecision:
        if not contract.safe_for_preserve:
            return TemplateApplicationDecision(
                mode=TemplateApplicationMode.PROFILE_ONLY,
                reason="active or external template content is blocked",
                numbering_owner=NumberingOwner.RENDERER,
                fidelity_score=contract.fidelity_score,
                diagnostics=contract.diagnostics,
            )
        owner = (
            NumberingOwner.TEMPLATE
            if contract.numbering.heading_owned
            else NumberingOwner.RENDERER
        )
        if explicit_mode is not None:
            return TemplateApplicationDecision(
                mode=explicit_mode,
                reason="user-explicit",
                numbering_owner=owner,
                fidelity_score=contract.fidelity_score,
                diagnostics=contract.diagnostics,
            )
        if contract.fidelity_score >= 0.8 and contract.slots:
            mode = TemplateApplicationMode.PRESERVE
            reason = "safe package with semantic slots"
        elif contract.fidelity_score >= 0.55 and contract.semantic_style_map:
            mode = TemplateApplicationMode.REMAP
            reason = "styles are reusable but package fidelity is partial"
        elif contract.fidelity_score < 0.35:
            mode = TemplateApplicationMode.CLARIFY
            reason = "template intent cannot be inferred reliably"
        else:
            mode = TemplateApplicationMode.PROFILE_ONLY
            reason = "reuse page and style profile only"
        return TemplateApplicationDecision(
            mode=mode,
            reason=reason,
            confirmation_required=mode is TemplateApplicationMode.CLARIFY,
            numbering_owner=owner,
            fidelity_score=contract.fidelity_score,
            diagnostics=contract.diagnostics,
        )


class TemplateRepairService:
    """Apply a safe template decision as a style/numbering-only DocumentIR revision."""

    def apply(
        self,
        document: DocumentIR,
        contract: TemplateContractV2,
        *,
        explicit_mode: TemplateApplicationMode | None = None,
    ) -> tuple[DocumentIR, TemplateApplicationDecision]:
        from paperagent.agents.document_ir import DocumentIR

        source = DocumentIR.model_validate(document)
        decision = TemplateContractResolver().resolve(contract, explicit_mode=explicit_mode)
        payload = source.model_dump(mode="json")
        metadata = dict(source.metadata)
        metadata.update(
            {
                "template_contract_hash": contract.contract_hash,
                "template_source_hash": contract.source_hash,
                "template_application_mode": decision.mode.value,
            }
        )
        payload["metadata"] = metadata
        heading_owner = decision.numbering_owner
        numbering = source.numbering.model_copy(
            update={
                "headings": source.numbering.headings.model_copy(
                    update={
                        "owner": heading_owner,
                        "source": f"template:{contract.source_hash}",
                    }
                ),
                "decision_source": f"template:{decision.mode.value}",
            }
        )
        payload["numbering"] = numbering.model_dump(mode="json")
        payload["revision"] = source.revision + 1
        return DocumentIR.model_validate(payload), decision


class TemplateContractStore:
    def __init__(self, root: Path) -> None:
        self.root = root

    def save(self, contract: TemplateContractV2) -> Path:
        self.root.mkdir(parents=True, exist_ok=True)
        path = self.root / f"{contract.source_hash}.json"
        path.write_text(contract.model_dump_json(indent=2), encoding="utf-8")
        return path

    def load(self, source_hash: str) -> TemplateContractV2 | None:
        path = self.root / f"{source_hash}.json"
        if not path.is_file():
            return None
        return TemplateContractV2.model_validate_json(path.read_text(encoding="utf-8"))


def migrate_template_contract(
    payload: dict[str, object], *, source_path: Path | None = None
) -> TemplateContractV2:
    if str(payload.get("schema_version", "")) == "2.0":
        return TemplateContractV2.model_validate(payload)
    if source_path is not None and source_path.is_file():
        return DocxTemplateParser().parse_contract(source_path)
    profile = TemplateProfile.model_validate(payload)
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
    diagnostic = TemplateDiagnostic(
        code="TEMPLATE_REEXTRACTION_RECOMMENDED",
        message="legacy profile lacks package-level fidelity; re-upload source for V2 extraction",
        severity=TemplateDiagnosticSeverity.WARNING,
    )
    return TemplateContractV2(
        source_hash=hashlib.sha256(canonical.encode("utf-8")).hexdigest(),
        page_sections=profile.section_pages or [profile.page],
        styles=profile.styles,
        semantic_style_map=profile.semantic_style_map,
        numbering=TemplateNumberingContract(
            owner=NumberingOwner.NONE,
            list_owned=profile.has_numbering,
            diagnostics=[diagnostic],
        ),
        theme_fonts=profile.theme_fonts,
        headers_footers=[
            item for item in profile.regions if item.kind in {"header", "footer"}
        ],
        table_styles=profile.table_styles,
        slots=[],
        fixed_content=[],
        capabilities=profile.capabilities,
        fidelity_score=0.5 if profile.semantic_style_map else 0.25,
        diagnostics=[diagnostic],
    )


class TemplateParseError(ValueError):
    pass


def points(value: object) -> float | None:
    return float(value.pt) if value is not None and hasattr(value, "pt") else None


class DocxTemplateParser:
    def parse(self, path: Path, *, completed_sample: bool = False) -> TemplateProfile:
        if not path.is_file():
            raise TemplateParseError(f"template does not exist: {path.name}")
        try:
            document = Document(str(path))
        except (OSError, ValueError, BadZipFile) as error:
            raise TemplateParseError(f"DOCX template cannot be parsed: {path.name}") from error
        first = document.sections[0]
        styles: list[StyleProfile] = []
        for style in document.styles:
            if style.type is None:
                continue
            element = style.element
            rpr = element.rPr
            east_asia = None
            ascii_font = None
            hansi_font = None
            if rpr is not None and rpr.rFonts is not None:
                east_asia = rpr.rFonts.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                )
                ascii_font = rpr.rFonts.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii"
                )
                hansi_font = rpr.rFonts.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi"
                )
            font = getattr(style, "font", None)
            paragraph_format = getattr(style, "paragraph_format", None)
            base_style = getattr(style, "base_style", None)
            raw_line_spacing = (
                paragraph_format.line_spacing if paragraph_format is not None else None
            )
            line_spacing = (
                float(raw_line_spacing)
                if isinstance(raw_line_spacing, (float, int))
                else points(raw_line_spacing)
            )
            styles.append(
                StyleProfile(
                    name=style.name,
                    style_type=str(style.type),
                    font_name=font.name if font else None,
                    east_asia_font=east_asia,
                    ascii_font=ascii_font,
                    hansi_font=hansi_font,
                    font_size_pt=points(font.size) if font else None,
                    bold=font.bold if font else None,
                    italic=font.italic if font else None,
                    alignment=(
                        str(paragraph_format.alignment)
                        if paragraph_format is not None
                        and paragraph_format.alignment is not None
                        else None
                    ),
                    line_spacing=line_spacing,
                    space_before_pt=(
                        points(paragraph_format.space_before)
                        if paragraph_format is not None
                        else None
                    ),
                    space_after_pt=(
                        points(paragraph_format.space_after)
                        if paragraph_format is not None
                        else None
                    ),
                    keep_with_next=(
                        paragraph_format.keep_with_next
                        if paragraph_format is not None
                        else None
                    ),
                    keep_together=(
                        paragraph_format.keep_together
                        if paragraph_format is not None
                        else None
                    ),
                    page_break_before=(
                        paragraph_format.page_break_before
                        if paragraph_format is not None
                        else None
                    ),
                    based_on=base_style.name if base_style else None,
                )
            )
        regions: list[TemplateRegion] = []
        for index, section in enumerate(document.sections):
            if section.header.paragraphs:
                text = "\n".join(item.text for item in section.header.paragraphs).strip()
                if text:
                    regions.append(TemplateRegion(kind="header", text=text, section_index=index))
            if section.footer.paragraphs:
                text = "\n".join(item.text for item in section.footer.paragraphs).strip()
                if text:
                    regions.append(TemplateRegion(kind="footer", text=text, section_index=index))
        headings = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style
            and paragraph.style.name.startswith("Heading")
            and paragraph.text.strip()
        ]
        first_page = [item.text for item in document.paragraphs[:8] if item.text.strip()]
        if first_page and not completed_sample:
            regions.append(TemplateRegion(kind="cover_candidate", text="\n".join(first_page)))
        numbering = bool(document.part.numbering_part.element.num_lst)
        semantic_styles = {
            item.name: semantic.value
            for item in styles
            if (semantic := semantic_style_id(item.name)) is not None
        }
        table_styles = sorted(item.name for item in styles if "TABLE" in item.style_type.upper())
        has_toc, numbering_count, theme_fonts = self._package_metadata(path)
        capabilities = ["page", "styles"]
        if regions:
            capabilities.append("header-footer")
        if numbering:
            capabilities.append("numbering")
        if has_toc:
            capabilities.append("toc")
        missing = [
            capability
            for capability in ("semantic-styles", "numbering")
            if (capability == "semantic-styles" and not semantic_styles)
            or (capability == "numbering" and not numbering)
        ]
        return TemplateProfile(
            source_name=path.name,
            page=PageProfile(
                width_pt=points(first.page_width),
                height_pt=points(first.page_height),
                top_margin_pt=points(first.top_margin),
                bottom_margin_pt=points(first.bottom_margin),
                left_margin_pt=points(first.left_margin),
                right_margin_pt=points(first.right_margin),
            ),
            section_pages=[
                PageProfile(
                    width_pt=points(section.page_width),
                    height_pt=points(section.page_height),
                    top_margin_pt=points(section.top_margin),
                    bottom_margin_pt=points(section.bottom_margin),
                    left_margin_pt=points(section.left_margin),
                    right_margin_pt=points(section.right_margin),
                )
                for section in document.sections
            ],
            section_count=len(document.sections),
            styles=styles,
            regions=regions,
            heading_structure=headings,
            has_numbering=numbering,
            numbering_abstract_count=numbering_count,
            has_toc=has_toc,
            table_styles=table_styles,
            theme_fonts=theme_fonts,
            semantic_style_map=semantic_styles,
            capabilities=capabilities,
            missing_capabilities=missing,
            completed_sample=completed_sample,
            warnings=(["完成样例仅提取结构和样式, 正文不进入新文档"] if completed_sample else [])
            + (["模板缺少语义样式映射, 需要确认默认样式"] if not semantic_styles else []),
        )

    def parse_contract(
        self, path: Path, *, completed_sample: bool = False
    ) -> TemplateContractV2:
        """Extract an immutable, non-executable contract from a DOCX package."""

        profile = self.parse(path, completed_sample=completed_sample)
        source_hash = hashlib.sha256(path.read_bytes()).hexdigest()
        diagnostics = self._security_diagnostics(path)
        numbering = self._numbering_contract(path)
        slots, fixed = self._slots_and_fixed(path, completed_sample=completed_sample)
        capabilities = set(profile.capabilities)
        if slots:
            capabilities.add("semantic-slots")
        if numbering.heading_owned:
            capabilities.add("heading-numbering")
        if profile.table_styles:
            capabilities.add("table-styles")
        score = 0.2
        score += 0.2 if profile.semantic_style_map else 0
        score += 0.15 if slots else 0
        score += 0.15 if profile.section_pages else 0
        score += 0.1 if profile.regions else 0
        score += 0.1 if profile.table_styles else 0
        score += 0.1 if numbering.heading_owned or numbering.list_owned else 0
        if any(item.severity is TemplateDiagnosticSeverity.ERROR for item in diagnostics):
            score = min(score, 0.3)
        return TemplateContractV2(
            source_hash=source_hash,
            page_sections=profile.section_pages,
            styles=profile.styles,
            semantic_style_map=profile.semantic_style_map,
            numbering=numbering,
            theme_fonts=profile.theme_fonts,
            headers_footers=[
                item for item in profile.regions if item.kind in {"header", "footer"}
            ],
            table_styles=profile.table_styles,
            slots=slots,
            fixed_content=fixed,
            capabilities=sorted(capabilities),
            fidelity_score=min(1, score),
            diagnostics=diagnostics,
        )

    @staticmethod
    def _security_diagnostics(path: Path) -> list[TemplateDiagnostic]:
        diagnostics: list[TemplateDiagnostic] = []
        with ZipFile(path) as archive:
            names = set(archive.namelist())
            if any(name.casefold().endswith("vbaproject.bin") for name in names):
                diagnostics.append(
                    TemplateDiagnostic(
                        code="TEMPLATE_MACRO_BLOCKED",
                        message="VBA project is never executed or copied to output",
                        severity=TemplateDiagnosticSeverity.ERROR,
                    )
                )
            if any(
                name.startswith(("word/embeddings/", "word/activeX/")) for name in names
            ):
                diagnostics.append(
                    TemplateDiagnostic(
                        code="TEMPLATE_EMBEDDED_OBJECT_BLOCKED",
                        message="embedded active objects are isolated",
                        severity=TemplateDiagnosticSeverity.ERROR,
                    )
                )
            for name in sorted(item for item in names if item.endswith(".rels")):
                payload = archive.read(name)
                if b'TargetMode="External"' in payload or b"TargetMode='External'" in payload:
                    diagnostics.append(
                        TemplateDiagnostic(
                            code="TEMPLATE_EXTERNAL_RELATIONSHIP_BLOCKED",
                            message=f"external relationship blocked in {name}",
                            severity=TemplateDiagnosticSeverity.ERROR,
                        )
                    )
        return diagnostics

    @staticmethod
    def _numbering_contract(path: Path) -> TemplateNumberingContract:
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        levels: list[TemplateNumberingLevel] = []
        diagnostics: list[TemplateDiagnostic] = []
        with ZipFile(path) as archive:
            if "word/numbering.xml" not in archive.namelist():
                return TemplateNumberingContract()
            root = ElementTree.fromstring(archive.read("word/numbering.xml"))
            for abstract in root.findall("w:abstractNum", namespace):
                abstract_id = abstract.attrib.get(f"{{{namespace['w']}}}abstractNumId", "0")
                for level in abstract.findall("w:lvl", namespace):
                    ilvl = int(level.attrib.get(f"{{{namespace['w']}}}ilvl", "0"))

                    def value(
                        name: str, default: str = "", *, _level: ElementTree.Element = level
                    ) -> str:
                        node = _level.find(f"w:{name}", namespace)
                        return (
                            node.attrib.get(f"{{{namespace['w']}}}val", default)
                            if node is not None
                            else default
                        )

                    levels.append(
                        TemplateNumberingLevel(
                            abstract_id=abstract_id,
                            level=ilvl,
                            format=value("numFmt", "decimal"),
                            level_text=value("lvlText", f"%{ilvl + 1}"),
                            start=int(value("start", "1") or "1"),
                            restart=(
                                int(value("lvlRestart")) if value("lvlRestart") else None
                            ),
                            paragraph_style=value("pStyle") or None,
                        )
                    )
        heading_owned = any(
            (item.paragraph_style or "").casefold().startswith("heading") for item in levels
        )
        list_owned = bool(levels)
        if list_owned and not heading_owned:
            diagnostics.append(
                TemplateDiagnostic(
                    code="TEMPLATE_LIST_NUMBERING_ONLY",
                    message="template numbering does not own semantic heading styles",
                    severity=TemplateDiagnosticSeverity.INFO,
                )
            )
        return TemplateNumberingContract(
            owner=NumberingOwner.TEMPLATE if heading_owned else NumberingOwner.NONE,
            heading_owned=heading_owned,
            list_owned=list_owned,
            levels=levels,
            style_links={
                item.paragraph_style: item.abstract_id
                for item in levels
                if item.paragraph_style
            },
            diagnostics=diagnostics,
        )

    @staticmethod
    def _slots_and_fixed(
        path: Path, *, completed_sample: bool
    ) -> tuple[list[TemplateSlot], list[TemplateFixedContent]]:
        document = Document(str(path))
        slots: list[TemplateSlot] = []
        fixed: list[TemplateFixedContent] = []
        labels = {
            "姓名": "author",
            "作者": "author",
            "班级": "class_name",
            "学校": "institution",
            "学院": "department",
            "指导教师": "advisor",
            "辅导老师": "advisor",
            "日期": "date",
            "题目": "title",
        }
        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name if paragraph.style else None
            matched_key = next((key for key in labels if key in text), None)
            placeholder = bool(
                re.search(
                    r"[_\uff3f]{2,}|[<\uff1c\u3010\[]"
                    r"[^>\uff1e\u3011\]]+[>\uff1e\u3011\]]",
                    text,
                )
            )
            if matched_key and (placeholder or index < 12):
                slots.append(
                    TemplateSlot(
                        slot_id=f"paragraph-{index}",
                        semantic_key=labels[matched_key],
                        label=matched_key,
                        location=f"paragraphs[{index}]",
                        required=matched_key in {"姓名", "作者", "题目"},
                        source_style=style,
                    )
                )
                continue
            category = "sample-content" if completed_sample else "fixed-label"
            fixed.append(
                TemplateFixedContent(
                    node_id=f"paragraph-{index}",
                    text_hash=hashlib.sha256(text.encode("utf-8")).hexdigest(),
                    category=category,
                    location=f"paragraphs[{index}]",
                    source_style=style,
                    reusable=not completed_sample and len(text) <= 30,
                )
            )
        return slots, fixed

    @staticmethod
    def _package_metadata(path: Path) -> tuple[bool, int, list[str]]:
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        theme_namespace = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        with ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
            numbering_count = 0
            if "word/numbering.xml" in archive.namelist():
                numbering = ElementTree.fromstring(archive.read("word/numbering.xml"))
                numbering_count = len(numbering.findall("w:abstractNum", namespace))
            theme_fonts: list[str] = []
            if "word/theme/theme1.xml" in archive.namelist():
                theme = ElementTree.fromstring(archive.read("word/theme/theme1.xml"))
                for item in theme.findall(".//a:latin", theme_namespace):
                    value = item.attrib.get("typeface")
                    if value and value not in theme_fonts:
                        theme_fonts.append(value)
            return (
                "TOC \\" in document_xml or "TOC \\o" in document_xml,
                numbering_count,
                theme_fonts,
            )
