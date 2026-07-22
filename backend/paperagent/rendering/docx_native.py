from __future__ import annotations

import re
from pathlib import Path
from typing import Any
from uuid import UUID

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image

from paperagent.agents.document_ir import (
    BlockKind,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    InlineKind,
    InlineNode,
    ListItem,
)
from paperagent.rendering.layout import (
    ArchetypeId,
    NamedStyle,
    Orientation,
    PageSpec,
    StyleProperties,
    archetype_layout_profile,
)
from paperagent.rendering.markdown_parser import math_aware_inline_nodes
from paperagent.rendering.math_markup import math_fragments
from paperagent.rendering.presentation_view import (
    RenderLine,
    RenderPresentationViewModel,
    RenderToken,
    RenderTokenKind,
)
from paperagent.rendering.template import (
    DocxTemplateParser,
    TemplateApplicationMode,
    TemplateContractResolver,
)
from paperagent.schemas.typography import TypographySpec

STYLE_NAMES: dict[NamedStyle, str] = {
    NamedStyle.DOCUMENT_TITLE: "Title",
    NamedStyle.SUBTITLE: "Subtitle",
    NamedStyle.HEADING_1: "Heading 1",
    NamedStyle.HEADING_2: "Heading 2",
    NamedStyle.HEADING_3: "Heading 3",
    NamedStyle.HEADING_4: "Heading 4",
    NamedStyle.HEADING_5: "Heading 5",
    NamedStyle.HEADING_6: "Heading 6",
    NamedStyle.BODY_TEXT: "BodyText",
    NamedStyle.LIST: "PaperAgent List",
    NamedStyle.QUOTE: "PaperAgent Quote",
    NamedStyle.CAPTION: "Caption",
    NamedStyle.TABLE: "Table Grid",
    NamedStyle.CODE: "PaperAgent Code",
    NamedStyle.EQUATION: "PaperAgent Equation",
    NamedStyle.REFERENCE: "PaperAgent Reference",
    NamedStyle.HEADER: "Header",
    NamedStyle.FOOTER: "Footer",
}


class NativeDocxRenderer:
    """DocumentIR 2.0 to editable OOXML with named styles and real fields."""

    def __init__(self) -> None:
        self._bookmark_id = 1
        self._template_mode = False

    def render(
        self,
        document: DocumentIR,
        output: Path,
        *,
        template: Path | None = None,
    ) -> Path:
        effective_template = template
        if template is not None:
            contract = DocxTemplateParser().parse_contract(template)
            decision = TemplateContractResolver().resolve(contract)
            if decision.mode in {
                TemplateApplicationMode.PROFILE_ONLY,
                TemplateApplicationMode.CLARIFY,
            }:
                effective_template = None
        self._template_mode = effective_template is not None
        word = self._new_document(effective_template)
        profile = archetype_layout_profile(
            self._archetype(document),
            language=document.language,
            explicit_theme=str(document.metadata.get("theme_id") or "") or None,
            project_theme=str(document.metadata.get("project_theme_id") or "") or None,
        )
        if effective_template is None:
            self._apply_page(word.sections[0], profile.page)
        self._install_styles(word, document, profile.styles.styles)
        self._configure_settings(word)
        presentation = RenderPresentationViewModel.from_document(document)
        self._configure_header_footer(word, presentation)
        self._cover(word, presentation)
        if profile.toc.enabled(
            estimated_pages=max(1, len(list(document.iter_blocks())) // 6),
            section_count=len(list(document.iter_sections())),
        ) or bool(document.metadata.get("toc")):
            self._toc(word)
        list_ids = self._numbering_definitions(word)
        for section in document.sections:
            self._section(word, section, document, list_ids, depth=1)
        references = self._references(document)
        if references:
            heading = "参考文献" if document.language in {"zh", "mixed"} else "References"
            word.add_heading(heading, level=1)
            for item in references:
                paragraph = word.add_paragraph(style=STYLE_NAMES[NamedStyle.REFERENCE])
                paragraph.add_run(item)
        output.parent.mkdir(parents=True, exist_ok=True)
        word.save(str(output))
        reopened = Document(str(output))
        if not reopened.paragraphs:
            raise ValueError("DOCX package contains no document paragraphs")
        return output

    @staticmethod
    def _new_document(template: Path | None) -> DocumentObject:
        word = Document(str(template)) if template else Document()
        if template:
            body = word._element.body
            for child in list(body):
                if child.tag != qn("w:sectPr"):
                    body.remove(child)
        return word

    @staticmethod
    def _archetype(document: DocumentIR) -> ArchetypeId:
        raw = document.metadata.get("archetype", ArchetypeId.RESEARCH_REPORT.value)
        try:
            return ArchetypeId(str(raw))
        except ValueError:
            return ArchetypeId.RESEARCH_REPORT

    @staticmethod
    def _apply_page(section: Any, page: PageSpec) -> None:
        width = page.width.points()
        height = page.height.points()
        if page.orientation is Orientation.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            width, height = height, width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Pt(width)
        section.page_height = Pt(height)
        section.top_margin = Pt(page.top_margin.points())
        section.bottom_margin = Pt(page.bottom_margin.points())
        section.left_margin = Pt(page.left_margin.points())
        section.right_margin = Pt(page.right_margin.points())
        section.gutter = Pt(page.gutter.points())
        section.different_first_page_header_footer = page.different_first_page
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            sect_pr.append(cols)
        cols.set(qn("w:num"), str(page.columns))

    def _install_styles(
        self,
        word: DocumentObject,
        document: DocumentIR,
        properties: dict[NamedStyle, StyleProperties],
    ) -> None:
        for semantic, name in STYLE_NAMES.items():
            if name not in word.styles:
                style_type = (
                    WD_STYLE_TYPE.TABLE if semantic is NamedStyle.TABLE else WD_STYLE_TYPE.PARAGRAPH
                )
                word.styles.add_style(name, style_type)
            if semantic is NamedStyle.TABLE:
                continue
            style = word.styles[name]
            spec = properties[semantic].model_copy(deep=True)
            if semantic is NamedStyle.BODY_TEXT:
                if document.typography.body_font:
                    spec.font_family = document.typography.body_font
                    spec.east_asia_font = document.typography.body_font
                if document.typography.body_size_pt:
                    spec.font_size_pt = document.typography.body_size_pt
                if document.typography.line_spacing:
                    spec.line_spacing = document.typography.line_spacing
            elif semantic.value.startswith("Heading"):
                if document.typography.heading_font:
                    spec.font_family = document.typography.heading_font
                    spec.east_asia_font = document.typography.heading_font
                if document.typography.heading_size_pt:
                    spec.font_size_pt = document.typography.heading_size_pt
            elif semantic is NamedStyle.CODE and document.typography.code_font:
                spec.font_family = document.typography.code_font
            self._apply_style(style, spec)
            if semantic is NamedStyle.BODY_TEXT:
                self._apply_style(word.styles["Normal"], spec)

    @staticmethod
    def _apply_style(style: Any, spec: StyleProperties) -> None:
        style.font.name = spec.font_family
        style.font.size = Pt(spec.font_size_pt)
        style.font.bold = spec.bold
        style.font.italic = spec.italic
        style.font.color.rgb = RGBColor.from_string(spec.color.removeprefix("#"))
        rpr = style._element.get_or_add_rPr()
        fonts = rpr.get_or_add_rFonts()
        fonts.set(qn("w:ascii"), spec.font_family)
        fonts.set(qn("w:hAnsi"), spec.font_family)
        fonts.set(qn("w:eastAsia"), spec.east_asia_font)
        paragraph = style.paragraph_format
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[spec.alignment]
        paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.line_spacing = spec.line_spacing
        paragraph.space_before = Pt(spec.space_before_pt)
        paragraph.space_after = Pt(spec.space_after_pt)
        paragraph.first_line_indent = Pt(
            spec.first_line_indent.points(font_size_pt=spec.font_size_pt)
        )
        paragraph.left_indent = Pt(spec.left_indent.points(font_size_pt=spec.font_size_pt))
        paragraph.right_indent = Pt(spec.right_indent.points(font_size_pt=spec.font_size_pt))
        paragraph.keep_with_next = spec.keep_with_next
        paragraph.keep_together = spec.keep_together
        paragraph.page_break_before = spec.page_break_before
        paragraph.widow_control = spec.widow_orphan
        ppr = style._element.get_or_add_pPr()
        border = ppr.find(qn("w:pBdr"))
        if border is not None:
            ppr.remove(border)
        if spec.shading:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), spec.shading.removeprefix("#"))
            style._element.get_or_add_pPr().append(shd)

    @staticmethod
    def _configure_settings(word: DocumentObject) -> None:
        settings = word.settings.element
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")
        compat = settings.find(qn("w:compat"))
        if compat is None:
            compat = OxmlElement("w:compat")
            settings.append(compat)

    def _configure_header_footer(
        self,
        word: DocumentObject,
        presentation: RenderPresentationViewModel,
    ) -> None:
        settings = word.settings.element
        odd_even = settings.find(qn("w:evenAndOddHeaders"))
        if presentation.different_odd_even and odd_even is None:
            odd_even = OxmlElement("w:evenAndOddHeaders")
            settings.append(odd_even)
        elif not presentation.different_odd_even and odd_even is not None:
            settings.remove(odd_even)
        for section in word.sections:
            section.different_first_page_header_footer = presentation.different_first_page
            default = (
                presentation.odd_page
                if presentation.different_odd_even
                else presentation.default
            )
            self._chrome_container(section.header, default.header, is_header=True)
            self._chrome_container(section.footer, default.footer, is_header=False)
            if presentation.different_first_page:
                self._chrome_container(
                    section.first_page_header,
                    presentation.first_page.header,
                    is_header=True,
                )
                self._chrome_container(
                    section.first_page_footer,
                    presentation.first_page.footer,
                    is_header=False,
                )
            if presentation.different_odd_even:
                self._chrome_container(
                    section.even_page_header,
                    presentation.even_page.header,
                    is_header=True,
                )
                self._chrome_container(
                    section.even_page_footer,
                    presentation.even_page.footer,
                    is_header=False,
                )

    def _chrome_container(self, container: Any, line: RenderLine, *, is_header: bool) -> None:
        if self._template_mode and (
            any(paragraph.text.strip() for paragraph in container.paragraphs)
            or bool(container.tables)
        ):
            return
        element = container._element
        for child in list(element):
            element.remove(child)
        base = OxmlElement("w:p")
        element.append(base)
        if not any((line.left, line.center, line.right)):
            return
        table = container.add_table(rows=1, cols=3, width=Inches(6.25))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for cell, tokens, alignment in zip(
            table.rows[0].cells,
            (line.left, line.center, line.right),
            (
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT,
            ),
            strict=True,
        ):
            cell.width = Inches(6.25 / 3)
            paragraph = cell.paragraphs[0]
            paragraph.style = STYLE_NAMES[
                NamedStyle.HEADER if is_header else NamedStyle.FOOTER
            ]
            paragraph.alignment = alignment
            for token in tokens:
                self._chrome_token(paragraph, token)
        self._remove_table_borders(table)

    def _chrome_token(self, paragraph: Any, token: RenderToken) -> None:
        if token.kind is RenderTokenKind.TEXT:
            paragraph.add_run(token.text)
        elif token.kind is RenderTokenKind.PAGE_NUMBER:
            self._field(paragraph, "PAGE")
        elif token.kind is RenderTokenKind.TOTAL_PAGES:
            self._field(paragraph, "NUMPAGES")
        else:
            self._field(paragraph, 'STYLEREF "Heading 1"', "")

    @staticmethod
    def _remove_table_borders(table: Any) -> None:
        properties = table._tbl.tblPr
        borders = properties.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            properties.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "nil")
            borders.append(node)

    @staticmethod
    def _field(paragraph: Any, instruction: str, placeholder: str = "1") -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        text = OxmlElement("w:instrText")
        text.set(qn("xml:space"), "preserve")
        text.text = f" {instruction} "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        value = OxmlElement("w:t")
        value.text = placeholder
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, text, separate, value, end))

    def _cover(
        self,
        word: DocumentObject,
        presentation: RenderPresentationViewModel,
    ) -> None:
        cover = presentation.cover
        if not cover.enabled:
            return
        title = word.add_paragraph(style=STYLE_NAMES[NamedStyle.DOCUMENT_TITLE])
        title.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[cover.alignment]
        title.paragraph_format.space_after = Pt(cover.title_spacing_after_pt)
        title.add_run(cover.title)
        last = title
        if cover.subtitle:
            last = word.add_paragraph(
                cover.subtitle,
                style=STYLE_NAMES[NamedStyle.SUBTITLE],
            )
        if cover.fields:
            table = word.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Mm(cover.max_content_width_mm * 0.32)
            table.columns[1].width = Mm(cover.max_content_width_mm * 0.68)
            for item in cover.fields:
                row = table.add_row()
                row.height = Pt(max(18, cover.field_row_spacing_pt + 11))
                row.cells[0].width = table.columns[0].width
                row.cells[1].width = table.columns[1].width
                row.cells[0].text = item.label
                row.cells[1].text = item.value
                for index, cell in enumerate(row.cells):
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.style = STYLE_NAMES[NamedStyle.BODY_TEXT]
                    paragraph.paragraph_format.keep_together = True
                    paragraph.paragraph_format.space_after = Pt(cover.field_row_spacing_pt)
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    )
                row_properties = row._tr.get_or_add_trPr()
                row_properties.append(OxmlElement("w:cantSplit"))
            self._remove_table_borders(table)
            last = word.add_paragraph()
            last.paragraph_format.space_before = Pt(0)
            last.paragraph_format.space_after = Pt(0)
        if cover.start_new_page_after:
            last.add_run().add_break(WD_BREAK.PAGE)

    def _toc(self, word: DocumentObject) -> None:
        heading = word.add_paragraph("目录", style=STYLE_NAMES[NamedStyle.HEADING_1])
        heading.paragraph_format.page_break_before = False
        paragraph = word.add_paragraph()
        self._field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "右键更新目录")
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def _section(
        self,
        word: DocumentObject,
        section: DocumentSection,
        document: DocumentIR,
        list_ids: dict[str, int],
        *,
        depth: int,
    ) -> None:
        level = max(1, min(6, section.level or depth))
        heading = word.add_heading(section.title, level=level)
        self._local_typography(
            heading,
            document.resolve_typography(section_id=section.section_id),
            BlockKind.HEADING,
        )
        self._bookmark(heading, section.section_id.hex)
        for block in section.blocks:
            self._block(word, block, document, list_ids, section.section_id)
        for child in section.children:
            self._section(word, child, document, list_ids, depth=depth + 1)

    def _block(
        self,
        word: DocumentObject,
        block: DocumentBlock,
        document: DocumentIR,
        list_ids: dict[str, int],
        section_id: UUID,
    ) -> None:
        typography = document.resolve_typography(
            section_id=section_id,
            block_id=block.block_id,
        )
        if block.kind is BlockKind.TABLE and block.table:
            self._table(word, block, typography, language=document.language)
            return
        if block.kind is BlockKind.FIGURE and block.figure and block.figure.path:
            self._figure(word, block, language=document.language)
            return
        if block.kind is BlockKind.LIST and block.list_spec:
            self._list(
                word,
                block.list_spec.items,
                list_ids,
                ordered=block.list_spec.kind.value == "ordered",
                typography=typography,
            )
            return
        if block.kind is BlockKind.EQUATION:
            self._equation(word, block, typography)
            return
        if block.kind is BlockKind.PAGE_BREAK:
            word.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            return
        if block.kind is BlockKind.SECTION_BREAK:
            created = word.add_section(WD_SECTION.NEW_PAGE)
            self._apply_page(created, PageSpec())
            self._configure_header_footer(
                word,
                RenderPresentationViewModel.from_document(document),
            )
            return
        if block.kind is BlockKind.HEADING:
            paragraph = word.add_heading(
                block.text,
                level=max(1, min(6, int(str(block.data.get("level", 2))))),
            )
        else:
            style = {
                BlockKind.PARAGRAPH: NamedStyle.BODY_TEXT,
                BlockKind.QUOTE: NamedStyle.QUOTE,
                BlockKind.CODE: NamedStyle.CODE,
                BlockKind.CITATION: NamedStyle.REFERENCE,
                BlockKind.CALLOUT: NamedStyle.QUOTE,
            }.get(block.kind, NamedStyle.BODY_TEXT)
            paragraph = word.add_paragraph(style=STYLE_NAMES[style])
            self._inlines(
                paragraph,
                block.inlines,
                block.text,
                typography=typography,
            )
            for number in self._citation_numbers(block, document):
                paragraph.add_run(f" [{number}]")
        self._bookmark(paragraph, block.block_id.hex)
        self._local_typography(paragraph, typography, block.kind)

    def _inlines(
        self,
        paragraph: Any,
        nodes: list[InlineNode],
        fallback: str,
        *,
        typography: TypographySpec,
    ) -> None:
        nodes = math_aware_inline_nodes(fallback, nodes)
        if not nodes:
            self._text_with_math(paragraph, fallback, typography)
            return
        for node in nodes:
            text = node.text or "".join(child.text for child in node.children)
            if node.kind is InlineKind.TEXT:
                self._text_with_math(paragraph, text, typography)
                continue
            if node.kind is InlineKind.LINK and node.href:
                self._hyperlink(paragraph, text, node.href)
                continue
            run = paragraph.add_run(text)
            run.bold = node.kind is InlineKind.STRONG
            run.italic = node.kind is InlineKind.EMPHASIS
            if node.kind is InlineKind.CODE:
                run.font.name = "Cascadia Mono"
            if node.kind is InlineKind.FOOTNOTE:
                run.font.superscript = True

    def _text_with_math(
        self, paragraph: Any, value: str, typography: TypographySpec
    ) -> None:
        previous = ""
        for kind, content in math_fragments(value):
            if kind == "text":
                paragraph.add_run(content)
                previous = content
                continue
            if kind == "display" and previous and not previous.endswith("\n"):
                paragraph.add_run().add_break()
            paragraph._p.append(
                self._native_math(
                    content.strip(),
                    mathml=None,
                    typography=typography,
                )
            )
            if kind == "display":
                paragraph.add_run().add_break()
            previous = ""

    @staticmethod
    def _hyperlink(paragraph: Any, text: str, url: str) -> None:
        relationship = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship)
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.extend((color, underline))
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.extend((run_properties, text_node))
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _list(
        self,
        word: DocumentObject,
        items: list[ListItem],
        list_ids: dict[str, int],
        *,
        ordered: bool,
        typography: TypographySpec,
        level: int = 0,
    ) -> None:
        for item in items:
            paragraph = word.add_paragraph(style=STYLE_NAMES[NamedStyle.LIST])
            self._numbered_paragraph(paragraph, list_ids["ordered" if ordered else "bullet"], level)
            self._inlines(
                paragraph,
                item.inlines,
                item.text,
                typography=typography,
            )
            self._local_typography(paragraph, typography, BlockKind.LIST)
            self._bookmark(paragraph, item.item_id.hex)
            if item.children:
                self._list(
                    word,
                    item.children,
                    list_ids,
                    ordered=False,
                    typography=typography,
                    level=min(level + 1, 8),
                )

    @staticmethod
    def _numbering_definitions(word: DocumentObject) -> dict[str, int]:
        numbering = word.part.numbering_part.element
        existing = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
        next_id = max(existing, default=0) + 1
        result: dict[str, int] = {}
        for kind in ("bullet", "ordered"):
            abstract_id = next_id
            abstract = OxmlElement("w:abstractNum")
            abstract.set(qn("w:abstractNumId"), str(abstract_id))
            multi = OxmlElement("w:multiLevelType")
            multi.set(qn("w:val"), "multilevel")
            abstract.append(multi)
            for level in range(9):
                lvl = OxmlElement("w:lvl")
                lvl.set(qn("w:ilvl"), str(level))
                start = OxmlElement("w:start")
                start.set(qn("w:val"), "1")
                num_fmt = OxmlElement("w:numFmt")
                num_fmt.set(qn("w:val"), "decimal" if kind == "ordered" else "bullet")
                lvl_text = OxmlElement("w:lvlText")
                lvl_text.set(qn("w:val"), f"%{level + 1}." if kind == "ordered" else "•")
                run_properties = OxmlElement("w:rPr")
                run_fonts = OxmlElement("w:rFonts")
                run_fonts.set(qn("w:ascii"), "Arial")
                run_fonts.set(qn("w:hAnsi"), "Arial")
                run_properties.append(run_fonts)
                ppr = OxmlElement("w:pPr")
                tabs = OxmlElement("w:tabs")
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), "num")
                tab.set(qn("w:pos"), str(720 + level * 360))
                tabs.append(tab)
                indent = OxmlElement("w:ind")
                indent.set(qn("w:left"), str(720 + level * 360))
                indent.set(qn("w:hanging"), "360")
                ppr.extend((tabs, indent))
                lvl.extend((start, num_fmt, lvl_text, ppr, run_properties))
                abstract.append(lvl)
            numbering.append(abstract)
            num = OxmlElement("w:num")
            num.set(qn("w:numId"), str(next_id))
            abstract_ref = OxmlElement("w:abstractNumId")
            abstract_ref.set(qn("w:val"), str(abstract_id))
            num.append(abstract_ref)
            numbering.append(num)
            result[kind] = next_id
            next_id += 1
        return result

    @staticmethod
    def _numbered_paragraph(paragraph: Any, num_id: int, level: int) -> None:
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num))
        ppr.append(num_pr)

    def _table(
        self,
        word: DocumentObject,
        block: DocumentBlock,
        typography: TypographySpec,
        *,
        language: str,
    ) -> None:
        assert block.table is not None
        rows = block.table.rows
        columns = max(len(row.cells) for row in rows)
        if block.caption:
            caption = word.add_paragraph(style=STYLE_NAMES[NamedStyle.CAPTION])
            self._sequence(caption, "Table", block.caption, language=language)
        table = word.add_table(rows=len(rows), cols=columns)
        table.style = (
            block.table.style_name
            if block.table.style_name in {item.name for item in word.styles}
            else STYLE_NAMES[NamedStyle.TABLE]
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._table_geometry(table, rows, columns, block.table.column_widths)
        for row_index, row in enumerate(rows):
            target_row = table.rows[row_index]
            cannot_split = OxmlElement("w:cantSplit")
            target_row._tr.get_or_add_trPr().append(cannot_split)
            if row_index == 0 and block.table.repeat_header:
                repeat = OxmlElement("w:tblHeader")
                repeat.set(qn("w:val"), "true")
                target_row._tr.get_or_add_trPr().append(repeat)
            for column, cell in enumerate(row.cells):
                target = target_row.cells[column]
                target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                target.text = ""
                self._inlines(
                    target.paragraphs[0],
                    cell.inlines,
                    cell.text,
                    typography=typography,
                )
                self._local_typography(
                    target.paragraphs[0],
                    typography,
                    BlockKind.TABLE,
                )
                if row_index == 0 or cell.header:
                    for run in target.paragraphs[0].runs:
                        run.bold = True

    @staticmethod
    def _table_geometry(
        table: Any,
        rows: list[Any],
        columns: int,
        configured_widths: list[float] | None = None,
    ) -> None:
        usable_dxa = int(PageSpec().content_width_pt * 20)
        weights = configured_widths or [
            max(
                4,
                min(
                    48,
                    max(
                        (len(row.cells[index].text) if index < len(row.cells) else 0)
                        for row in rows
                    ),
                ),
            )
            for index in range(columns)
        ]
        total = sum(weights)
        widths = [max(720, round(usable_dxa * weight / total)) for weight in weights]
        difference = usable_dxa - sum(widths)
        widths[-1] += difference
        tbl_pr = table._tbl.tblPr
        table_width = tbl_pr.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            tbl_pr.append(table_width)
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(usable_dxa))
        indent = tbl_pr.find(qn("w:tblInd"))
        if indent is None:
            indent = OxmlElement("w:tblInd")
            tbl_pr.append(indent)
        indent.set(qn("w:type"), "dxa")
        indent.set(qn("w:w"), "0")
        margins = tbl_pr.find(qn("w:tblCellMar"))
        if margins is None:
            margins = OxmlElement("w:tblCellMar")
            tbl_pr.append(margins)
        for side in ("top", "left", "bottom", "right"):
            value = margins.find(qn(f"w:{side}"))
            if value is None:
                value = OxmlElement(f"w:{side}")
                margins.append(value)
            value.set(qn("w:w"), "100" if side in {"top", "bottom"} else "120")
            value.set(qn("w:type"), "dxa")
        grid = table._tbl.tblGrid
        for grid_col, width in zip(grid.gridCol_lst, widths, strict=True):
            grid_col.set(qn("w:w"), str(width))
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                cell.width = Pt(widths[index] / 20)
                cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                cell_width.set(qn("w:type"), "dxa")
                cell_width.set(qn("w:w"), str(widths[index]))

    def _figure(self, word: DocumentObject, block: DocumentBlock, *, language: str) -> None:
        assert block.figure is not None and block.figure.path is not None
        path = Path(block.figure.path)
        if not path.is_file():
            return
        paragraph = word.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Keep the image with its following caption. The shared Caption style
        # intentionally keeps a table caption with the table that follows, but
        # a figure caption comes after its image and must not attach itself to
        # the next section heading.
        paragraph.paragraph_format.keep_with_next = True
        max_width = 6.2 * (block.figure.width_ratio or 1)
        max_height = 8.25
        try:
            with Image.open(path) as image:
                width_px, height_px = image.size
            if width_px > 0 and height_px > 0:
                max_width = min(max_width, max_height * width_px / height_px)
        except OSError:
            # SVG and other validated formats are converted to PNG before DOCX rendering.
            pass
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(max_width))
        drawing = run._r.xpath(".//wp:docPr")
        if drawing:
            drawing[0].set("descr", block.figure.alt_text or block.caption or "Figure")
            drawing[0].set("title", block.caption or "Figure")
        self._bookmark(paragraph, block.block_id.hex)
        caption = word.add_paragraph(style=STYLE_NAMES[NamedStyle.CAPTION])
        caption.paragraph_format.keep_with_next = False
        self._sequence(caption, "Figure", block.caption or "", language=language)

    def _sequence(
        self,
        paragraph: Any,
        sequence: str,
        caption: str,
        *,
        language: str,
    ) -> None:
        label = {"Table": "表", "Figure": "图"}.get(sequence, sequence)
        paragraph.add_run(f"{label if language in {'zh', 'mixed'} else sequence} ")
        self._field(paragraph, f"SEQ {sequence} \\* ARABIC", "1")
        body = caption_body(caption, sequence)
        if body:
            paragraph.add_run(f"  {body}")

    def _equation(
        self,
        word: DocumentObject,
        block: DocumentBlock,
        typography: TypographySpec,
    ) -> None:
        latex = block.equation.latex if block.equation else block.text
        paragraph = word.add_paragraph(style=STYLE_NAMES[NamedStyle.EQUATION])
        self._local_typography(paragraph, typography, BlockKind.EQUATION)
        math_para = OxmlElement("m:oMathPara")
        mathml = block.equation.mathml if block.equation else None
        math_para.append(self._native_math(latex, mathml=mathml, typography=typography))
        paragraph._p.append(math_para)
        self._bookmark(paragraph, block.block_id.hex)

    @staticmethod
    def _native_math(
        latex: str,
        *,
        mathml: str | None,
        typography: TypographySpec,
    ) -> Any:
        """Create editable OMML instead of placing LaTeX source inside an m:t node."""

        import mathml2omml
        from latex2mathml.converter import convert as latex_to_mathml

        source = mathml or latex_to_mathml(latex)
        omml = mathml2omml.convert(source)
        namespace = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        if "xmlns:m=" not in omml:
            omml = omml.replace("<m:oMath", f'<m:oMath xmlns:m="{namespace}"', 1)
        math = parse_xml(omml)
        font = typography.equation_font or typography.body_font
        size = typography.equation_size_pt or typography.body_size_pt
        for run in math.iter(qn("m:r")):
            run_properties = run.find(qn("w:rPr"))
            if run_properties is None:
                run_properties = OxmlElement("w:rPr")
                math_properties = run.find(qn("m:rPr"))
                insert_at = 1 if math_properties is not None else 0
                run.insert(insert_at, run_properties)
            if font:
                fonts = OxmlElement("w:rFonts")
                for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                    fonts.set(qn(f"w:{attribute}"), font)
                run_properties.append(fonts)
            if size:
                for name in ("w:sz", "w:szCs"):
                    value = OxmlElement(name)
                    value.set(qn("w:val"), str(round(size * 2)))
                    run_properties.append(value)
        return math

    @staticmethod
    def _local_typography(
        paragraph: Any,
        typography: TypographySpec,
        kind: BlockKind,
    ) -> None:
        if kind is BlockKind.HEADING:
            font = typography.heading_font or typography.body_font
            size = typography.heading_size_pt or typography.body_size_pt
        elif kind is BlockKind.TABLE:
            font = typography.table_font or typography.body_font
            size = typography.table_size_pt or typography.body_size_pt
        elif kind is BlockKind.CODE:
            font = typography.code_font or typography.body_font
            size = typography.code_size_pt or typography.body_size_pt
        elif kind is BlockKind.EQUATION:
            font = typography.equation_font or typography.body_font
            size = typography.equation_size_pt or typography.body_size_pt
        else:
            font = typography.body_font
            size = typography.body_size_pt
        if typography.line_spacing:
            paragraph.paragraph_format.line_spacing = typography.line_spacing
        for run in paragraph.runs:
            if font:
                run.font.name = font
                fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                fonts.set(qn("w:ascii"), font)
                fonts.set(qn("w:hAnsi"), font)
                fonts.set(qn("w:eastAsia"), font)
            if size:
                run.font.size = Pt(size)

    def _bookmark(self, paragraph: Any, identity: str) -> None:
        name = "pa_" + re.sub(r"[^A-Za-z0-9_]", "", identity)[:35]
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)
        self._bookmark_id += 1

    @staticmethod
    def _citation_numbers(block: DocumentBlock, document: DocumentIR) -> list[int]:
        manifest = document.metadata.get("evidence_manifest", [])
        if not isinstance(manifest, list):
            return []
        by_id = {
            str(item.get("evidence_id")): index
            for index, item in enumerate(manifest, start=1)
            if isinstance(item, dict)
        }
        return sorted(
            {
                by_id[str(item.evidence_id)]
                for item in block.citations
                if str(item.evidence_id) in by_id
            }
        )

    @staticmethod
    def _references(document: DocumentIR) -> list[str]:
        manifest = document.metadata.get("evidence_manifest", [])
        if not isinstance(manifest, list):
            return []
        seen: set[str] = set()
        references: list[str] = []
        for item in manifest:
            if not isinstance(item, dict):
                continue
            source = str(item.get("source_uri", ""))
            key = source or str(item.get("title", ""))
            if not key or key in seen:
                continue
            seen.add(key)
            references.append(
                f"[{len(references) + 1}] {item.get('title', 'Untitled')}"
                + (f" — {source}" if source else "")
            )
        return references


def caption_body(value: str, sequence: str) -> str:
    """Remove an author-supplied label because Word supplies the SEQ label/number."""

    if not value.strip():
        return ""
    labels = ("Table", "表") if sequence == "Table" else ("Figure", "图")
    punctuation = ":\uFF1A.\u3001-"
    label_pattern = "|".join(re.escape(label) for label in labels)
    pattern = rf"^\s*(?:{label_pattern})\s*\d*\s*[{punctuation}]?\s*"
    return re.sub(pattern, "", value, count=1, flags=re.IGNORECASE).strip()
