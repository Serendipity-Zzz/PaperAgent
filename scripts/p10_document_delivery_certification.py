from __future__ import annotations

import asyncio
import json
import os
import re
import time
from datetime import UTC, datetime
from pathlib import Path
from typing import cast
from uuid import UUID, uuid4
from zipfile import ZipFile

import psutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from pydantic import JsonValue
from pypdf import PdfReader, PdfWriter

from paperagent.agents.change_intent import ChangeIntent, ChangeScope
from paperagent.agents.document_ir import (
    BlockKind,
    CitationRef,
    DocumentBlock,
    DocumentIR,
    DocumentSection,
    EquationSpec,
    FigureSpec,
    FrontMatter,
    Provenance,
    TableCell,
    TableRow,
    TableSpec,
)
from paperagent.artifacts import ArtifactService, CompletionClaimValidator
from paperagent.core.config import Settings
from paperagent.db import DatabaseManager
from paperagent.execution.tool_suite import ExecutionToolSuite
from paperagent.preview import PreviewService
from paperagent.providers import ChatMessage, ChatRequest, ProviderModality
from paperagent.providers.adapters import OpenAICompatibleProvider
from paperagent.providers.registry import ProviderRegistry
from paperagent.rendering import (
    BibliographicItem,
    CitationStyle,
    CitationStyleService,
    DocumentRevisionStore,
    RevisionOperation,
    RevisionWorkflow,
)
from paperagent.schemas.typography import TypographySpec
from paperagent.security.credentials import CredentialStore

PRIVATE_PLACEHOLDER = "Verified source content is supplied by the renderer."


def _field(paragraph: Paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, end])


def _template(path: Path) -> None:
    word = Document()
    section = word.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    header = section.header.paragraphs[0]
    header.text = "PaperAgent | 实验报告模板"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    _field(footer, "PAGE")
    footer.add_run(" 页")
    for name, size in (("Normal", 11), ("Heading 1", 16), ("Heading 2", 14)):
        style = word.styles[name]
        style.font.name = "SimSun" if name == "Normal" else "SimHei"
        style.font.size = Pt(size)
    word.add_heading("模板示例正文(渲染时必须移除)", level=1)
    word.add_paragraph("This sample body must never appear in generated output.")
    path.parent.mkdir(parents=True, exist_ok=True)
    word.save(str(path))


def _experiment_source() -> str:
    return """from pathlib import Path
import csv
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

length = 1.0
frequency = 60.0
wave_speed = 120.0
amplitude = 0.02
x = np.linspace(0.0, length, 501)
k = 2.0 * np.pi * frequency / wave_speed
y = 2.0 * amplitude * np.sin(k * x)

with Path("standing_wave.csv").open("w", encoding="utf-8", newline="") as stream:
    writer = csv.writer(stream)
    writer.writerow(["x_m", "displacement_m"])
    writer.writerows(zip(x, y, strict=True))

with Path("standing_wave_large.csv").open("w", encoding="utf-8", newline="") as stream:
    writer = csv.writer(stream)
    writer.writerow(["sample", "x_m", "displacement_m"])
    for index in range(10000):
        position = length * index / 9999
        writer.writerow([index, position, 2.0 * amplitude * np.sin(k * position)])

fig, axis = plt.subplots(figsize=(8.0, 4.5), layout="constrained")
axis.plot(x, y, color="#339CFF", linewidth=2.2, label="t = 0 s")
axis.axhline(0.0, color="#555555", linewidth=0.8)
axis.set(xlabel="Position x (m)", ylabel="Displacement y (m)", title="Standing-wave displacement")
axis.grid(alpha=0.25)
axis.legend()
fig.savefig("standing_wave.png", dpi=180)
fig.savefig("standing_wave.svg")
"""


async def _real_provider_probe(settings: Settings) -> dict[str, object]:
    databases = DatabaseManager(settings)
    databases.initialize_global()
    registry = ProviderRegistry(databases)
    config = registry.active(ProviderModality.TEXT, project_id=None) or next(
        (item for item in registry.list() if item.credential_ref), None
    )
    if config is None or not config.credential_ref:
        raise RuntimeError("no encrypted real Provider is configured")
    credentials = CredentialStore(settings.resolved_data_dir / "global" / "credentials.json")
    provider = OpenAICompatibleProvider(
        config,
        lambda: credentials.get(config.credential_ref) if config.credential_ref else None,
    )
    started = time.perf_counter()
    try:
        response = await provider.chat(
            ChatRequest(
                messages=[
                    ChatMessage(
                        role="user",
                        content=(
                            "Call document_structure_probe exactly once for an experimental "
                            "report. Do not answer with prose."
                        ),
                    )
                ],
                max_tokens=256,
                tools=[
                    {
                        "type": "function",
                        "function": {
                            "name": "document_structure_probe",
                            "description": "Select the document archetype for certification.",
                            "parameters": {
                                "type": "object",
                                "properties": {
                                    "archetype": {"type": "string"},
                                    "requires_experiment": {"type": "boolean"},
                                },
                                "required": ["archetype", "requires_experiment"],
                                "additionalProperties": False,
                            },
                        },
                    }
                ],
            )
        )
        if not response.tool_calls or response.tool_calls[0].name != "document_structure_probe":
            raise RuntimeError("real Provider did not select the required registered tool")
        return {
            "provider_id": config.id,
            "provider_type": config.provider_type,
            "model": response.model,
            "tool_call": True,
            "latency_ms": round((time.perf_counter() - started) * 1000),
        }
    finally:
        await provider.client.aclose()


def _artifact_id(payload: JsonValue) -> str:
    if not isinstance(payload, dict):
        raise AssertionError("tool output did not include a registered artifact id")
    artifact_id = payload.get("artifact_id")
    if not isinstance(artifact_id, str):
        raise AssertionError("tool output did not include a registered artifact id")
    return artifact_id


def _document(figure_id: str, source_id: UUID) -> DocumentIR:
    provenance = Provenance(agent="p10_certification", evidence_ids=[source_id])
    citation = CitationRef(evidence_id=source_id, locator="Sec. 3", verified=True)
    bibliography = BibliographicItem(
        title="Vibrations and Waves",
        authors=["A. P. French"],
        year=1971,
        publisher="W. W. Norton",
        item_type="book",
        verified=True,
        source_evidence_id=source_id,
    )
    apa = CitationStyleService().format(bibliography, CitationStyle.APA, sequence=1)
    return DocumentIR(
        requirement_id=uuid4(),
        requirement_version=1,
        outline_id=uuid4(),
        title="弦上驻波特性实验报告",
        language="mixed",
        front_matter=FrontMatter(
            subtitle="真实 Python 实验与多格式排版认证",
            authors=["PaperAgent"],
            organization="Local-first Document Laboratory",
            date="2026-07-18",
            abstract=(
                "本报告在受控本地 uv 环境中运行 Python,计算一米长弦在 60 Hz "
                "激励下的驻波位移,并以同一份结构化文档生成 Markdown、DOCX 与 PDF。"
            ),
            keywords=["驻波", "本地实验", "DocumentIR", "XeLaTeX"],
        ),
        typography=TypographySpec(
            body_font="SimSun",
            heading_font="SimHei",
            table_font="SimSun",
            equation_font="Cambria Math",
            body_size_pt=11,
            heading_size_pt=16,
            table_size_pt=10.5,
            line_spacing=1.5,
            first_line_indent_chars=2,
        ),
        sections=[
            DocumentSection(
                title="实验背景与理论",
                goal="说明驻波条件与理论模型",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text=(
                            "驻波由频率相同、传播方向相反的两列波叠加形成。弦长、张力与"
                            "线密度共同决定允许的本征频率;波节处位移恒为零,波腹处振幅最大。"
                        ),
                        citations=[citation],
                        provenance=provenance,
                    ),
                    DocumentBlock(
                        kind=BlockKind.EQUATION,
                        text=r"y(x,t)=2A\sin(kx)\cos(\omega t)",
                        equation=EquationSpec(
                            latex=r"y(x,t)=2A\sin(kx)\cos(\omega t)",
                            number=True,
                            label="eq:standing-wave",
                        ),
                        caption="驻波位移方程",
                        provenance=provenance,
                    ),
                ],
            ),
            DocumentSection(
                title="实验设计与参数",
                goal="记录可复现实验设置",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text=(
                            "程序在无网络的托管运行目录中执行。采样使用 501 个等距位置点,"
                            "并额外生成 10000 行数据用于大表格预览性能验证。"
                        ),
                        provenance=provenance,
                    ),
                    DocumentBlock(
                        kind=BlockKind.TABLE,
                        caption="表 1 驻波实验参数",
                        table=TableSpec(
                            rows=[
                                TableRow(
                                    cells=[
                                        TableCell(text="参数", header=True),
                                        TableCell(text="值", header=True),
                                        TableCell(text="单位", header=True),
                                    ]
                                ),
                                TableRow(
                                    cells=[
                                        TableCell(text="弦长"),
                                        TableCell(text="1.0"),
                                        TableCell(text="m"),
                                    ]
                                ),
                                TableRow(
                                    cells=[
                                        TableCell(text="频率"),
                                        TableCell(text="60"),
                                        TableCell(text="Hz"),
                                    ]
                                ),
                                TableRow(
                                    cells=[
                                        TableCell(text="波速"),
                                        TableCell(text="120"),
                                        TableCell(text="m/s"),
                                    ]
                                ),
                            ],
                            column_widths=[0.42, 0.28, 0.30],
                        ),
                        provenance=provenance,
                    ),
                ],
            ),
            DocumentSection(
                title="实验结果与讨论",
                goal="呈现真实数据图并解释节点分布",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text=(
                            "计算结果显示位移包络随位置呈正弦变化。图中曲线直接来自本轮"
                            "Python 执行产物,并通过 artifact hash 与 run 记录建立追溯关系。"
                        ),
                        provenance=provenance,
                    ),
                    DocumentBlock(
                        kind=BlockKind.FIGURE,
                        caption="图 1 弦上驻波位移分布",
                        figure=FigureSpec(
                            artifact_id=UUID(figure_id),
                            alt_text="一米长弦上的驻波位移曲线",
                            width_ratio=0.82,
                        ),
                        provenance=provenance,
                    ),
                ],
            ),
            DocumentSection(
                title="结论",
                goal="总结结果与证据边界",
                blocks=[
                    DocumentBlock(
                        kind=BlockKind.PARAGRAPH,
                        text=(
                            "实验完成了源码、数据、图像与文档的同轮交付。结果验证了驻波"
                            "节点与波腹的空间分布,同时证明三种输出格式可共享正文、公式、"
                            "表格、图片和引用,而无需把 Markdown 原文伪装成排版文件。"
                        ),
                        provenance=provenance,
                    )
                ],
            ),
        ],
        back_matter=[
            DocumentBlock(
                kind=BlockKind.CITATION,
                text=apa.bibliography,
                provenance=provenance,
            )
        ],
        metadata={
            "archetype": "experiment-report",
            "toc": True,
            "pdf_render_mode": "xelatex",
            "design_preset": "standard_business_brief",
            "design_overrides": ["A4", "SimSun/SimHei", "1.5 line spacing"],
            "evidence_manifest": [
                {
                    "title": bibliography.title,
                    "source_uri": "https://archive.org/details/vibrationswaves00fren",
                    "verified": True,
                }
            ],
        },
    )


def _long_preview_fixtures(root: Path, artifacts: ArtifactService, run_id: str) -> list[str]:
    directory = root / "artifacts" / "preview-stress"
    directory.mkdir(parents=True, exist_ok=True)
    pdf = directory / "fifty-five-pages.pdf"
    writer = PdfWriter()
    for _ in range(55):
        writer.add_blank_page(width=595, height=842)
    with pdf.open("wb") as stream:
        writer.write(stream)
    markdown = directory / "long-report.md"
    markdown.write_text(
        "# Long preview\n\n"
        + "\n\n".join(f"## Section {i}\n\nContent {i}." for i in range(1, 501)),
        encoding="utf-8",
    )
    return [
        artifacts.register(pdf, kind="output", producer_tool="p10.preview", run_id=run_id).id,
        artifacts.register(markdown, kind="output", producer_tool="p10.preview", run_id=run_id).id,
    ]


def _assert_native_documents(paths: list[Path]) -> None:
    for path in paths:
        if not path.is_file() or path.stat().st_size == 0:
            raise AssertionError(f"missing output: {path.name}")
        if path.suffix == ".pdf":
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            if not reader.pages or PRIVATE_PLACEHOLDER in text:
                raise AssertionError(f"invalid PDF: {path.name}")
            if not any(page.images for page in reader.pages):
                raise AssertionError(f"PDF embedded experiment image missing: {path.name}")
        elif path.suffix == ".docx":
            word = Document(str(path))
            text = "\n".join(paragraph.text for paragraph in word.paragraphs)
            if "This sample body must never appear" in text or PRIVATE_PLACEHOLDER in text:
                raise AssertionError(f"template or placeholder leaked: {path.name}")
            if not word.inline_shapes or not word.tables:
                raise AssertionError(f"DOCX native image/table missing: {path.name}")
        elif path.suffix == ".md":
            text = path.read_text(encoding="utf-8")
            if "![" not in text or "##" not in text or PRIVATE_PLACEHOLDER in text:
                raise AssertionError(f"Markdown structure/image missing: {path.name}")
            image_targets = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", text)
            if not image_targets or not all(
                (path.parent / target).resolve().is_file() for target in image_targets
            ):
                raise AssertionError(f"Markdown image asset missing: {path.name}")
        elif path.suffix == ".zip":
            with ZipFile(path) as bundle:
                names = set(bundle.namelist())
                if "report.md" not in names or not any(
                    name.startswith("assets/") and name.lower().endswith((".png", ".jpg", ".svg"))
                    for name in names
                ):
                    raise AssertionError(f"portable Markdown bundle incomplete: {path.name}")


async def main() -> int:
    workspace = Path(__file__).resolve().parents[1]
    root = Path(os.environ.get("PAPERAGENT_P10_CERT_ROOT", workspace / ".tmp" / "p10-k-cert"))
    root = root.resolve()
    root.mkdir(parents=True, exist_ok=True)
    settings = Settings(
        project_root=workspace,
        data_dir=root / "data",
        environment="test",
    )
    started = time.perf_counter()
    provider_evidence = await _real_provider_probe(Settings())

    databases = DatabaseManager(settings)
    databases.initialize_global()
    project_id = str(uuid4())
    project_root = databases.project_root(project_id)
    project_root.mkdir(parents=True, exist_ok=True)
    databases.project_engine(project_id).dispose()
    run_id = str(uuid4())
    artifacts = ArtifactService(databases, project_id)
    suite = ExecutionToolSuite(
        data_root=settings.resolved_data_dir,
        project_root=project_root,
        run_id=run_id,
        uv_path=Path(r"E:\App\uv\current\uv.exe"),
        artifact_service=artifacts,
        source_conversation_id="p10-k-certification",
    )
    try:
        machine = suite.machine_inspect({})
        environment = suite.environment_prepare(
            {
                "dependencies": ["matplotlib==3.10.3", "numpy==2.3.1"],
                "python_version": "3.12",
                "cuda_version": None,
            }
        )
        source = suite.code_materialize(
            {"filename": "standing_wave_experiment.py", "content": _experiment_source()}
        )
        execution = suite.process_execute(
            {"argv": ["python", "standing_wave_experiment.py"], "timeout_seconds": 180}
        )
        if not isinstance(execution, dict) or execution.get("status") != "completed":
            raise RuntimeError(f"real standing-wave experiment failed: {execution}")
        collected = suite.result_collect({})
        if not isinstance(collected, dict):
            raise AssertionError("result collection did not return artifact records")
        raw_artifacts = collected.get("artifacts")
        if not isinstance(raw_artifacts, list):
            raise AssertionError("result collection did not return artifact records")
        by_name = {
            str(item["name"]): item
            for item in raw_artifacts
            if isinstance(item, dict) and isinstance(item.get("name"), str)
        }
        required_experiment = {
            "standing_wave_experiment.py",
            "standing_wave.csv",
            "standing_wave_large.csv",
            "standing_wave.png",
            "standing_wave.svg",
        }
        if not required_experiment <= set(by_name):
            raise AssertionError(
                f"experiment outputs missing: {required_experiment - set(by_name)}"
            )
        figure_id = _artifact_id(cast(JsonValue, by_name["standing_wave.png"]))

        evidence_id = uuid4()
        document = _document(figure_id, evidence_id)
        render_payloads = [
            suite.document_render(
                {
                    "document_ir": cast(JsonValue, document.canonical_payload()),
                    "format": format_name,
                    "filename": f"standing-wave-report.{format_name}",
                    **({"pdf_mode": "xelatex"} if format_name == "pdf" else {}),
                }
            )
            for format_name in ("md", "md_bundle", "docx", "pdf")
        ]
        rendered_ids = [_artifact_id(payload) for payload in render_payloads]
        rendered = [artifacts.get(artifact_id) for artifact_id in rendered_ids]
        rendered_paths = [artifacts.verify(item) for item in rendered]
        _assert_native_documents(rendered_paths)
        CompletionClaimValidator(artifacts).validate(
            run_id,
            "已生成包含实验图的报告 standing-wave-report.md、DOCX 和 PDF",
        )

        store = DocumentRevisionStore(
            project_root,
            databases=databases,
            project_id=project_id,
            artifact_service=artifacts,
        )
        revision_one = store.load(document.document_id, 1)
        hashes_one = revision_one.hashes()
        workflow = RevisionWorkflow(store)
        global_revision = workflow.apply(
            revision_one,
            RevisionOperation(
                kind="typography",
                patch=ChangeIntent(
                    scope=ChangeScope.GLOBAL,
                    typography_patch={"body_font": "SimSun"},
                ).model_dump(mode="json"),
            ),
        ).document
        section_revision = workflow.apply(
            global_revision,
            RevisionOperation(
                kind="typography",
                patch=ChangeIntent(
                    scope=ChangeScope.SECTION,
                    section_ids=[global_revision.sections[1].section_id],
                    typography_patch={"heading_font": "SimHei", "heading_size_pt": 16},
                ).model_dump(mode="json"),
            ),
        ).document
        figure_block = next(
            block for block in section_revision.iter_blocks() if block.kind is BlockKind.FIGURE
        )
        revised = workflow.resize_figure(section_revision, figure_block.block_id, 0.70).document
        for candidate in (global_revision, section_revision, revised):
            hashes = candidate.hashes()
            if (
                hashes.content_hash != hashes_one.content_hash
                or hashes.structure_hash != hashes_one.structure_hash
                or hashes.asset_set_hash != hashes_one.asset_set_hash
                or hashes.citation_set_hash != hashes_one.citation_set_hash
            ):
                raise AssertionError("style-only revision changed semantic/resource hashes")
        if store.load(document.document_id, 1).hashes() != hashes_one:
            raise AssertionError("revision one was not preserved")

        template_path = project_root / "runs" / run_id / "experiment-template.docx"
        _template(template_path)
        template_artifact = artifacts.register(
            template_path,
            kind="template",
            producer_tool="user.upload",
            run_id=run_id,
        )
        template_docx = suite.document_render(
            {
                "document_id": str(revised.document_id),
                "revision": revised.revision,
                "format": "docx",
                "filename": "standing-wave-report-template.docx",
                "template_artifact_id": template_artifact.id,
            }
        )
        word_pdf = suite.document_render(
            {
                "document_id": str(revised.document_id),
                "revision": revised.revision,
                "format": "pdf",
                "filename": "standing-wave-report-word-parity.pdf",
                "pdf_mode": "word_parity",
                "template_artifact_id": template_artifact.id,
            }
        )
        revised_md = suite.document_render(
            {
                "document_id": str(revised.document_id),
                "revision": revised.revision,
                "format": "md",
                "filename": "standing-wave-report-revised.md",
            }
        )
        final_ids = [
            *rendered_ids,
            _artifact_id(template_docx),
            _artifact_id(word_pdf),
            _artifact_id(revised_md),
        ]
        final_artifacts = [artifacts.get(item) for item in final_ids]
        _assert_native_documents([artifacts.verify(item) for item in final_artifacts])

        citation_styles = {
            style.value: CitationStyleService()
            .format(
                BibliographicItem(
                    title="Vibrations and Waves",
                    authors=["A. P. French"],
                    year=1971,
                    publisher="W. W. Norton",
                    verified=True,
                ),
                style,
                sequence=1,
            )
            .model_dump(mode="json")
            for style in CitationStyle
        }
        if set(citation_styles) != {"gb-t-7714", "apa", "ieee"}:
            raise AssertionError("citation style parameterization is incomplete")

        stress_ids = _long_preview_fixtures(project_root, artifacts, run_id)
        preview_candidates = [
            *final_ids,
            _artifact_id(cast(JsonValue, by_name["standing_wave.png"])),
            _artifact_id(cast(JsonValue, by_name["standing_wave_experiment.py"])),
            _artifact_id(cast(JsonValue, by_name["standing_wave_large.csv"])),
            *stress_ids,
        ]
        preview = PreviewService(project_root)
        process = psutil.Process()
        rss_before = process.memory_info().rss
        preview_started = time.perf_counter()
        preview_records: list[dict[str, object]] = []
        try:
            for artifact_id in preview_candidates:
                artifact = artifacts.get(artifact_id)
                rendered_preview = preview.render(
                    artifacts.verify(artifact),
                    file_id=artifact.id,
                    source_hash=artifact.sha256,
                    source_name=artifact.original_name,
                )
                if rendered_preview.status.value not in {"ready", "failed"}:
                    raise AssertionError("preview did not reach a terminal cache state")
                preview_records.append(
                    {
                        "name": artifact.original_name,
                        "status": rendered_preview.status.value,
                        "fidelity": rendered_preview.fidelity.value,
                        "parts": rendered_preview.part_count,
                    }
                )
            first = artifacts.get(rendered_ids[-1])
            repeated = [
                preview.render(
                    artifacts.verify(first),
                    file_id=first.id,
                    source_hash=first.sha256,
                    source_name=first.original_name,
                ).id
                for _ in range(20)
            ]
            if len(set(repeated)) != 1:
                raise AssertionError(
                    "same preview opened repeatedly created duplicate cache entries"
                )
        finally:
            preview.close()
        preview_ms = round((time.perf_counter() - preview_started) * 1000)
        rss_delta_mb = round((process.memory_info().rss - rss_before) / (1024 * 1024), 2)

        public_artifacts = [artifacts.payload(item) for item in final_artifacts]
        if any(str(project_root) in json.dumps(item) for item in public_artifacts):
            raise AssertionError("public artifact payload leaked an absolute project path")
        report: dict[str, object] = {
            "status": "passed",
            "executed_at": datetime.now(UTC).isoformat(),
            "provider": provider_evidence,
            "machine": machine,
            "environment": {
                "fingerprint": environment.get("fingerprint")
                if isinstance(environment, dict)
                else None,
                "reused": environment.get("reused") if isinstance(environment, dict) else None,
            },
            "execution": {
                "status": execution.get("status") if isinstance(execution, dict) else None,
                "source_artifact_id": _artifact_id(source),
                "outputs": sorted(required_experiment),
                "simulated_data": False,
            },
            "document": {
                "document_id": str(revised.document_id),
                "revisions": [
                    1,
                    global_revision.revision,
                    section_revision.revision,
                    revised.revision,
                ],
                "formats": sorted(item.original_name for item in final_artifacts),
                "pdf_engines": [
                    payload.get("render_engine")
                    for payload in (render_payloads[-1], word_pdf)
                    if isinstance(payload, dict)
                ],
                "citation_styles": sorted(citation_styles),
                "semantic_hashes_preserved": True,
                "template_sample_removed": True,
                "embedded_images": True,
            },
            "preview": {
                "records": preview_records,
                "same_file_20x_single_cache_id": True,
                "elapsed_ms": preview_ms,
                "rss_delta_mb": rss_delta_mb,
            },
            "public_payload_absolute_path_leak": False,
            "latency_ms": round((time.perf_counter() - started) * 1000),
        }
        report_path = Path(
            os.environ.get(
                "PAPERAGENT_P10_CERT_REPORT",
                workspace / "docs" / "test-reports" / "document-layout" / "P10-K-live.json",
            )
        ).resolve()
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        manifest = root / "visual-manifest.json"
        manifest.write_text(
            json.dumps(
                {
                    "project_root": str(project_root),
                    "artifacts": [
                        {
                            "id": item.id,
                            "name": item.original_name,
                            "path": str(artifacts.verify(item)),
                            "sha256": item.sha256,
                        }
                        for item in final_artifacts
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        print(report_path)
        print(manifest)
        return 0
    finally:
        suite.close()


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
