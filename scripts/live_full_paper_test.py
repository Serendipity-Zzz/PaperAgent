from __future__ import annotations

import asyncio
import json
import os
import platform
import shutil
import time
from datetime import UTC, datetime
from pathlib import Path

import httpx
import matplotlib.pyplot as plt
from docx import Document

from paperagent.agents.change_intent import ChangeIntent, ChangeScope
from paperagent.agents.document_ir import BlockKind, DocumentBlock, Provenance
from paperagent.agents.evidence import (
    EvidenceItem,
    EvidenceKind,
    LiteratureEvidenceAgent,
    VerificationStatus,
)
from paperagent.agents.outline import OutlineDesignerAgent, TemplateSection
from paperagent.agents.requirements import RequirementUnderstandingAgent
from paperagent.agents.review import RepairPlanner, ReviewAgent, ReviewPolicy
from paperagent.agents.state import RawRequest, RequirementSpec
from paperagent.agents.writer import SectionWriterAgent
from paperagent.core.config import Settings
from paperagent.db.manager import DatabaseManager
from paperagent.engine import AgentLoop
from paperagent.experiments.runtime import (
    CapabilityAnalyzer,
    EnvironmentManager,
    EnvironmentRegistry,
    ExecutionApproval,
    ExperimentResultPackage,
    ProcessExecutor,
)
from paperagent.ingestion.classification import DocumentClassifier
from paperagent.ingestion.parsers import default_registry
from paperagent.ingestion.schemas import ImportReport
from paperagent.knowledge.models import Confidentiality, KnowledgeItem, KnowledgeScope
from paperagent.knowledge.retrieval import RetrievalResult
from paperagent.knowledge.service import ProjectKnowledgeService, report_to_items
from paperagent.literature import LiteratureRecord
from paperagent.orchestration import CandidatePlanGenerator, PlanValidator, safe_fallback_plan
from paperagent.providers import ModelProvider, ProviderConfig
from paperagent.providers.adapters import (
    AnthropicProvider,
    GeminiProvider,
    OpenAICompatibleProvider,
)
from paperagent.providers.registry import ProviderRegistry
from paperagent.providers.routing import ProviderRouter
from paperagent.rendering import TargetedTypographyService
from paperagent.security.credentials import CredentialStore
from paperagent.tools import ToolExecutor, ToolRegistry, ToolResultStore


def configured_provider(settings: Settings) -> tuple[ProviderConfig, ModelProvider]:
    databases = DatabaseManager(settings)
    databases.initialize_global()
    config = next(
        (item for item in ProviderRegistry(databases).list() if item.credential_ref), None
    )
    if config is None:
        raise RuntimeError("no configured text Provider credential")
    credentials = CredentialStore(settings.resolved_data_dir / "global" / "credentials.json")

    def credential() -> str | None:
        return credentials.get(config.credential_ref) if config.credential_ref else None

    provider = (
        AnthropicProvider(config, credential)
        if config.provider_type == "anthropic"
        else GeminiProvider(config, credential)
        if config.provider_type == "gemini"
        else OpenAICompatibleProvider(config, credential)
    )
    return config, provider


def download_sources(root: Path) -> tuple[Path, Path]:
    source_dir = root / "input"
    source_dir.mkdir(parents=True, exist_ok=True)
    pdf = source_dir / "react-arxiv-2210.03629.pdf"
    cached = next(
        (
            item
            for item in sorted(
                root.parent.glob("full-paper-*/input/react-arxiv-2210.03629.pdf"),
                key=lambda item: item.stat().st_mtime_ns,
                reverse=True,
            )
            if item.resolve() != pdf.resolve()
            and item.stat().st_size > 100_000
            and item.read_bytes()[:4] == b"%PDF"
        ),
        None,
    )
    if cached is not None:
        shutil.copy2(cached, pdf)
    else:
        failures: list[str] = []
        for url in (
            "https://arxiv.org/pdf/2210.03629",
            "https://export.arxiv.org/pdf/2210.03629",
        ):
            try:
                response = httpx.get(
                    url,
                    follow_redirects=True,
                    timeout=60,
                    headers={"User-Agent": "PaperAgent-local-validation/0.1"},
                )
                response.raise_for_status()
                if len(response.content) <= 100_000 or response.content[:4] != b"%PDF":
                    raise ValueError("downloaded arXiv payload is not a valid-sized PDF")
                pdf.write_bytes(response.content)
                break
            except (httpx.HTTPError, ValueError) as error:
                failures.append(f"{url}: {error.__class__.__name__}")
        else:
            raise RuntimeError(f"open literature download failed: {failures}")
    notes = source_dir / "experiment-notes.docx"
    word = Document()
    word.add_heading("本地智能体检索实验记录", level=1)
    word.add_paragraph(
        "目标: 测量 PaperAgent 本地 FTS 对 ReAct, reasoning, acting 和 智能体 "
        "关键词的真实查询延迟。"
    )
    word.add_paragraph("约束: 实验使用隔离 uv 环境, 禁用网络, 结果不得使用模拟数据。")
    word.add_paragraph("验收: 保存逐次延迟、均值、P95、环境指纹和可追溯数据图。")
    word.save(str(notes))
    return pdf, notes


def ingest_sources(
    root: Path, project_id: str, pdf: Path, notes: Path
) -> tuple[
    ProjectKnowledgeService,
    list[KnowledgeItem],
    list[ImportReport],
    list[RetrievalResult],
    list[RetrievalResult],
]:
    registry = default_registry()
    classifier = DocumentClassifier()
    service = ProjectKnowledgeService(root)
    selected = []
    reports = []
    for source, uri in (
        (pdf, "https://arxiv.org/abs/2210.03629"),
        (notes, None),
    ):
        report = registry.import_file(source)
        reports.append(report)
        classification = classifier.classify(report.source)
        items = report_to_items(
            report,
            classification,
            collection_id="live-full-paper",
            scope=KnowledgeScope.PROJECT,
            project_id=project_id,
            source_uri=uri,
            confidentiality=Confidentiality.PERSONAL,
        )
        service.ingest(items)
        selected.extend(items[:4])
    english = service.retrieve("ReAct reasoning acting", project_id=project_id, limit=5)
    chinese = service.retrieve("智能体 检索 实验", project_id=project_id, limit=5)
    if not english or not chinese:
        raise AssertionError("cross-language local retrieval did not return both query sets")
    locators = [item.locator.model_dump(exclude_none=True) for item in selected]
    if not any("page" in item for item in locators):
        raise AssertionError("PDF ingestion did not preserve page locators")
    if not any("paragraph" in item for item in locators):
        raise AssertionError("DOCX ingestion did not preserve paragraph locators")
    return service, selected, reports, english, chinese


def run_experiment(root: Path, knowledge_db: Path) -> tuple[ExperimentResultPackage, Path]:
    runtime_root = root / "runtimes"
    registry = EnvironmentRegistry(runtime_root)
    uv = Path("E:/App/uv/current/uv.exe").resolve()
    environment = EnvironmentManager(registry, uv, root / "uv-cache").ensure(
        "python-stdlib-only==3.12",
        python_version="3.12",
    )
    capability = CapabilityAnalyzer().analyze(root, gpu_name=None, vram_gb=None)
    run_dir = root / "experiment"
    run_dir.mkdir(parents=True, exist_ok=True)
    metrics_path = run_dir / "metrics.json"
    code = (
        "import json,sqlite3,statistics,sys,time;"
        "db,out=sys.argv[1:3];c=sqlite3.connect(db);samples=[];"
        "queries=['ReAct','reasoning','acting','智能体'];"
        "[(lambda s:(c.execute(\"SELECT count(*) FROM knowledge_fts WHERE "
        "knowledge_fts MATCH ?\",(q,)).fetchone(),samples.append((time.perf_counter()-s)"
        "*1000)))(time.perf_counter()) for _ in range(10) for q in queries];"
        "samples.sort();json.dump({'samples_ms':samples,'mean_ms':statistics.mean(samples),'p95_ms':samples[int(len(samples)*0.95)-1],'queries':queries},open(out,'w',encoding='utf-8'),ensure_ascii=False,indent=2)"
    )
    python = Path(environment.path) / "Scripts" / "python.exe"
    approval = ExecutionApproval(
        command=[str(python), "-c", code, str(knowledge_db), str(metrics_path)],
        working_directory=str(run_dir),
        writable_paths=[str(run_dir)],
        network_allowed=False,
        timeout_seconds=60,
        approved=True,
    )
    result = ProcessExecutor().run(approval)
    if result.status != "completed" or not metrics_path.is_file():
        raise RuntimeError(f"real retrieval experiment failed: {result.status} {result.stderr}")
    metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
    chart = run_dir / "retrieval-latency.png"
    plt.figure(figsize=(8, 4.5))
    plt.plot(metrics["samples_ms"], color="#339CFF", linewidth=1.5)
    plt.axhline(metrics["mean_ms"], color="#EF8354", linestyle="--", label="mean")
    plt.xlabel("Measured query sequence")
    plt.ylabel("Latency (ms)")
    plt.title("PaperAgent local FTS retrieval latency")
    plt.legend()
    plt.tight_layout()
    plt.savefig(chart, dpi=160)
    plt.close()
    package = ExperimentResultPackage(
        repository=str(root),
        commit="working-tree-live-validation",
        environment_fingerprint=environment.fingerprint,
        command=approval.command,
        hardware={
            "platform": platform.platform(),
            "capability_verdict": capability.verdict.value,
            "ram_gb": capability.ram_gb,
        },
        metrics={"mean_ms": metrics["mean_ms"], "p95_ms": metrics["p95_ms"]},
        data_files=[str(metrics_path)],
        figures=[str(chart)],
        status="completed",
        simulated_data=False,
    )
    if not package.eligible_as_experiment_evidence:
        raise AssertionError("real experiment package was not eligible as evidence")
    return package, chart


async def main() -> int:
    workspace = Path(__file__).resolve().parents[1]
    settings = Settings()
    config, provider = configured_provider(settings)
    run_id = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    root = workspace / "docs" / "test-reports" / "artifacts" / f"full-paper-{run_id}"
    root.mkdir(parents=True, exist_ok=False)
    started = time.perf_counter()
    try:
        pdf, notes = download_sources(root)
        project_id = "live-full-paper"
        knowledge, knowledge_items, reports, english, chinese = ingest_sources(
            root, project_id, pdf, notes
        )
        package, chart = run_experiment(root, knowledge.index.path)

        raw_text = (
            "写一篇约600个中文字符的项目报告, 主题是本地论文智能体的检索与闭环设计; "
            "面向软件工程研究者, 采用APA引用, 必须检索开放文献并运行已获授权的本地检索延迟实验, "
            "使用真实数据图, 不生成AI现场图, 输出Markdown、DOCX、Typst、LaTeX和PDF。"
        )
        requirement_agent = RequirementUnderstandingAgent(provider)
        understood = await requirement_agent.understand(RawRequest(text=raw_text))
        payload = understood.model_dump(mode="json")
        payload.update(
            status="draft",
            open_questions=[],
            conflicts=[],
            normalized_request=raw_text,
            research_formulation={
                "research_topic": "本地论文智能体的检索与闭环设计",
                "research_objective": "验证可追溯检索、真实实验和定向返修闭环",
                "scope_and_boundaries": ["个人 Windows 本地部署", "不包含图片生成"],
                "methodology_candidates": ["开放文献检索", "本地 FTS 延迟实验"],
                "data_requirements": ["真实查询延迟", "PDF/DOCX 页段定位"],
            },
            document_type="project_report",
            primary_language="zh",
            target_length={"value": 600, "unit": "chinese_char"},
            audience="软件工程研究者",
            citation_style="APA",
            requires_literature_search=True,
            requires_experiment=True,
            requires_data_chart=True,
            requires_generated_image=False,
            output_formats=["md", "docx", "typst", "latex", "pdf"],
            typography={"body_font": "SimSun", "body_size_pt": 11},
        )
        requirement = RequirementSpec.model_validate(payload).confirm()

        tool_registry = ToolRegistry()
        loop = AgentLoop(
            ProviderRouter([provider]),
            tool_registry,
            ToolExecutor(tool_registry, ToolResultStore(root / "tool-results")),
            root,
        )
        candidates = await CandidatePlanGenerator(loop).generate(
            requirement,
            project_id=project_id,
            available_tools=["knowledge.search", "experiment.run", "document.render"],
        )
        candidate = candidates.candidates[candidates.recommended_index]
        validation = PlanValidator(
            tool_registry,
            allowed_agents={item.agent_type for item in candidate.nodes},
        ).validate(candidate)
        selected_plan = candidate if validation.valid else safe_fallback_plan(requirement)

        initial_outline = OutlineDesignerAgent().design(requirement)
        outline = OutlineDesignerAgent.adjust(
            initial_outline,
            [
                TemplateSection(title="背景与证据", goal="说明问题和开放文献", weight=1),
                TemplateSection(title="实现与实验", goal="说明闭环、环境和真实结果", weight=2),
                TemplateSection(title="结论与局限", goal="总结证据边界", weight=1),
            ],
            "真实全链验收采用三节短报告以控制 Provider 成本",
        )
        literature = LiteratureRecord(
            title="ReAct: Synergizing Reasoning and Acting in Language Models",
            authors=("Shunyu Yao", "Jeffrey Zhao", "Dian Yu", "et al."),
            year=2022,
            doi=None,
            source="arxiv",
            source_uri="https://arxiv.org/abs/2210.03629",
            abstract=(
                "ReAct interleaves reasoning traces and task-specific actions, allowing language "
                "models to interact with external sources and update action plans."
            ),
            license="arXiv distribution license",
            open_access=True,
        )
        evidence = LiteratureEvidenceAgent().build(
            outline,
            literature=[literature],
            knowledge=knowledge_items,
        )
        experiment_evidence = EvidenceItem(
            kind=EvidenceKind.EXPERIMENT,
            title="PaperAgent local FTS latency experiment",
            content=json.dumps(package.metrics, ensure_ascii=False),
            source_id=str(package.run_id),
            verification=VerificationStatus.VERIFIED,
            locator={"metrics_file": package.data_files[0], "figure": package.figures[0]},
            reason="由获批隔离环境运行产生, 非模拟数据",
        )
        evidence = evidence.model_copy(update={"items": [*evidence.items, experiment_evidence]})
        policy = ReviewPolicy(level="standard", length_tolerance=0.5)
        preflight = ReviewAgent().preflight(requirement, outline, evidence, None, policy)
        if not preflight.passed:
            raise AssertionError("preflight review failed")
        writer = SectionWriterAgent(provider)
        drafts = [
            await writer.generate_section(requirement, section, evidence)
            for section in outline.sections
        ]
        document = writer.assemble(
            requirement,
            outline,
            evidence,
            drafts,
            title="本地论文智能体的检索与闭环设计验证",
        )
        document.sections[1].blocks.append(
            DocumentBlock(
                kind=BlockKind.FIGURE,
                caption="本地 FTS 真实检索延迟",
                data={"path": str(chart)},
                provenance=Provenance(
                    agent="experiment_agent",
                    evidence_ids=[experiment_evidence.evidence_id],
                ),
            )
        )
        postflight = ReviewAgent().postflight(document, requirement, evidence, policy)
        repair = RepairPlanner().plan(postflight.issues, round=1)

        renderer = TargetedTypographyService(root)
        first = renderer.apply(
            document,
            ChangeIntent(
                scope=ChangeScope.GLOBAL,
                typography_patch={"body_font": "SimSun", "body_size_pt": 11},
            ),
            formats=["md", "docx", "typst", "latex", "pdf"],
        )
        target_block = first.document.sections[0].blocks[0]
        revised = renderer.apply(
            first.document,
            ChangeIntent(
                scope=ChangeScope.BLOCK,
                block_ids=[target_block.block_id],
                typography_patch={"body_font": "SimSun", "body_size_pt": 12},
            ),
            formats=["md", "docx", "typst", "latex", "pdf"],
        )
        if revised.render_errors:
            raise RuntimeError(f"render errors: {revised.render_errors}")
        markdown = next(root / item.path for item in revised.artifacts if item.format == "md")
        markdown_text = markdown.read_text(encoding="utf-8")
        if "## References" not in markdown_text or "arxiv.org/abs/2210.03629" not in markdown_text:
            raise AssertionError("rendered references are not traceable")
        if revised.visual_diff is None:
            raise AssertionError("preview typography revision produced no visual diff")
        if revised.invalidation.affected_block_ids != [target_block.block_id]:
            raise AssertionError("typography revision invalidated unrelated content")

        report = {
            "status": "passed",
            "executed_at": datetime.now(UTC).isoformat(),
            "provider_id": config.id,
            "model": config.model,
            "requirement_prompt_hash": (
                requirement_agent.last_compiled_prompt.prompt_hash
                if requirement_agent.last_compiled_prompt
                else None
            ),
            "dynamic_plan": {
                "candidate_count": len(candidates.candidates),
                "recommended_valid": validation.valid,
                "selected_plan_id": str(selected_plan.plan_id),
                "selected_node_count": len(selected_plan.nodes),
            },
            "sources": [
                {
                    "name": report.source.name,
                    "parser": report.source.parser,
                    "chunks": len(report.source.chunks),
                }
                for report in reports
            ],
            "retrieval": {
                "english_hits": len(english),
                "chinese_hits": len(chinese),
                "english_locators": [item.hit.locator for item in english[:3]],
                "chinese_locators": [item.hit.locator for item in chinese[:3]],
            },
            "experiment": package.model_dump(mode="json"),
            "evidence_items": len(evidence.items),
            "reference_evidence_ids": [str(item) for item in evidence.reference_evidence_ids],
            "review": {
                "preflight_passed": preflight.passed,
                "postflight_passed": postflight.passed,
                "issues": [item.model_dump(mode="json") for item in postflight.issues],
                "repair_tasks": [item.model_dump(mode="json") for item in repair.tasks],
            },
            "output": {
                "document_id": str(revised.document.document_id),
                "revision": revised.document.revision,
                "formats": sorted(item.format for item in revised.artifacts),
                "artifacts": [item.model_dump(mode="json") for item in revised.artifacts],
                "visual_changed_pages": revised.visual_diff.changed_pages,
                "targeted_block_ids": [
                    str(item) for item in revised.invalidation.affected_block_ids
                ],
                "content_regenerated": revised.invalidation.regenerate_text,
            },
            "latency_ms": round((time.perf_counter() - started) * 1000),
        }
        target = Path(
            os.environ.get(
                "PAPERAGENT_LIVE_FULL_EVIDENCE",
                workspace / "docs" / "test-reports" / "P5-R-live-full-paper.json",
            )
        ).resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(target)
        return 0
    finally:
        client = getattr(provider, "client", None)
        if client is not None:
            await client.aclose()


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
